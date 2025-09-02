# === 0) FUTURE / TYPE CHECKS ===================================================
from __future__ import annotations

# === 1) STANDARD LIB IMPORTS ===================================================
import os, re, io, csv, json, time, uuid, difflib, hashlib, zipfile, base64
import threading
from io import BytesIO
from pathlib import Path
from collections import defaultdict, Counter
from datetime import date, datetime, timedelta, timezone
from typing import Optional, List, Dict, Any, Tuple, Literal, Callable, Iterable
from urllib.parse import urlparse
from difflib import SequenceMatcher

# === 2) THIRD-PARTY IMPORTS ====================================================
from dotenv import load_dotenv
from fastapi import FastAPI, Query, Header, HTTPException, Request, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse, Response, HTMLResponse
from starlette.responses import JSONResponse as StarletteJSONResponse, StreamingResponse as StarletteStreamingResponse
from pydantic import BaseModel, EmailStr, Field, confloat, conint
from pydantic_settings import BaseSettings, SettingsConfigDict
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# OCR / parsing
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except Exception:
    pdf_extract_text = None
try:
    import docx2txt
except Exception:
    docx2txt = None
try:
    import pytesseract
    import cv2
    from PIL import Image
except Exception:
    pytesseract = None
    cv2 = None
    Image = None

# === INTERNALS ================================================================
from security import ensure_auth, approval_gate, _guard_ip_from_request, APPROVAL_HEADER
from research_agent import (
    _perform_research,
    _effective_research_mode,
    research_extract,
    _extract_urls,
    kb_search_tfidf,
)




load_dotenv()
URL_RE = re.compile(r'(https?://[^\s)>\]]+)')

try:
    import docx2txt
except Exception:
    docx2txt = None

try:
    import cv2
    import pytesseract
except Exception:
    cv2 = None
    pytesseract = None

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except Exception:
    pdf_extract_text = None

try:
    RankItem  # type: ignore[name-defined]
except NameError:
    class RankItem(BaseModel):
        """
        Elemento di ranking generico (es: combinazione trigger/persona).
        """
        id: Optional[str] = None
        score: confloat = Field(0.0, ge=0.0, le=1.0)
        payload: Dict[str, Any] = Field(default_factory=dict)

try:
    CalendarEvent  # type: ignore[name-defined]
except NameError:
    class CalendarEvent(BaseModel):
        """
        Evento minimo usato per la calendarizzazione (fallback).
        """
        date: Optional[str] = None           # ISO (yyyy-mm-dd) oppure gestito altrove
        action: str = "send"
        channel: Optional[str] = None
        step: Optional[str] = None
        no_weekend_respected: bool = True
        meta: Dict[str, Any] = Field(default_factory=dict)

# === 3) FASTAPI APP (se non già esistente) ====================================

docs_on = os.getenv("DOCS_ENABLED", "true").lower() in ("1","true","yes","on")
APP_ENV = os.getenv("APP_ENV", "dev").lower()

app = FastAPI(
    title="Flowagent V3 Orchestrator",
    version="1.1.0",
    docs_url="/docs" if docs_on else None,
    redoc_url=None,
    openapi_url="/openapi.json",
)

# === 4) COSTANTI DI BASE / PATHS ==============================================
KB_DIR = os.getenv("KB_DIR", "./kb")
KB_RAW = Path(KB_DIR)
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "./uploads")
ALLOW_PER_CONTACT = os.getenv("ALLOW_PER_CONTACT", "false").lower() == "true"
APPROVAL_HEADER = "X-Connector-Approved"
IDEMPOTENCY_HEADER = "Idempotency-Key"
BEARER_TOKEN = os.getenv("BEARER_TOKEN", "")
preview_chars = 800
TRUTHY = {"true", "yes", "1", "on"}
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "./uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)


TAVILY_BUDGET: int = 900
SERPER_BUDGET: int = 2000
SERPAPI_BUDGET: int = 300

ALWAYS_ON = os.getenv("RESEARCH_ALWAYS_ON", "true").lower() in ("1","true","yes","on")

DEFAULT_TIMEOUT = (5, 15)

ALLOWED_CIDRS: str = "0.0.0.0/0"
LOG_LEVEL: str = "INFO"

PERSONA_DIR = os.getenv("PERSONA_DIR", "./kb/personas")
PERSONA_INDEX_JSON = os.getenv("PERSONA_INDEX_JSON", "./kb/persona_index.json")
WEBHOOK_SECRET = os.getenv("WEBHOOK_SECRET", "")
_PERSONA_INDEX: dict | None = None

class ComplianceRequest(BaseModel):
    text: str
    rules: Optional[Dict[str, Any]] = None

class ComplianceViolation(BaseModel):
    code: str
    detail: str

class ComplianceResponse(BaseModel):
    passed: bool = Field(alias="pass", default=True)
    violations: List[ComplianceViolation] = Field(default_factory=list)

    class Config:
        allow_population_by_field_name = True

model_config = {  # pydantic v2 (se sei già su v2)
        "env_file": ".env",
        "case_sensitive": True,
    }

class Settings(BaseSettings):
    OPENAI_API_KEY: str | None = None
    OPENAI_MODEL: str = "gpt-4o-mini"
    model_config = SettingsConfigDict(env_file=".env", case_sensitive=True)

settings = Settings()

RESEARCH_PROVIDER = os.getenv("RESEARCH_PROVIDER", "router").lower()   # serper | brave | ...
FALLBACKS = [p.strip().lower() for p in os.getenv("RESEARCH_FALLBACKS", "brave,tavily,serper,serpapi,searxng").split(",") if p.strip()]
DEGRADE_MODE = os.getenv("RESEARCH_DEGRADE_TO", "kb_only")  # "kb_only".strip()
SERPAPI_API_KEY   = os.getenv("SERPAPI_API_KEY")
TAVILY_API_KEY    = os.getenv("TAVILY_API_KEY")
SERPER_API_KEY=os.getenv("SERPER_API_KEY")
BRAVE_API_KEY=os.getenv("BRAVE_API_KEY")

raw_origins = os.getenv(
    "ALLOW_ORIGINS",
    "https://chatgpt.com,https://chat.openai.com,https://flowagent-v3-orchestrator.onrender.com"
)

ALLOWED_ORIGINS = [o.strip() for o in raw_origins.split(",") if o.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,            # richiede che NON ci sia "*"
    allow_methods=["*"],
    allow_headers=["*"],               # include Authorization, X-Connector-Approved
    expose_headers=["X-Connector-Approved", "Idempotency-Key"],
    max_age=3600,
)

# --- Messaggistica / Sequenze ------------------------
class Message(BaseModel):
    channel: Channel
    step: str
    variant: Literal["A", "B", "C"] = "A"
    subject: Optional[str] = None
    text: str
    tips: Optional[List[str]] = None

class COI(BaseModel):
    status: Literal["none", "estimated", "computed"] = "none"
    note: Optional[str] = None
    assumptions: Optional[List[str]] = None

class SequenceAction(BaseModel):
    channel: str
    step: str
    action: str
    date: Optional[date] = None
    id: Optional[str] = None
    meta: Optional[Dict[str, Any]] = None

# --- KB / Ricerca ----------------------------------------------

class SourceRef(BaseModel):
        title: Optional[str] = None
        url: Optional[str] = None
        snippet: Optional[str] = None
        published_at: Optional[str] = None

class ResearchResult(BaseModel):
    facts: Optional[List[EnrichedFact]] = None
    citations: Optional[List[Citation]] = None
    triggers: Optional[Dict[str, Any]] = None

# --- Personas -------------------------------------------------------
class PersonaDoc(BaseModel):
        persona_id: str
        industry: Optional[str] = None
        role: Optional[str] = None
        pains: List[str] = []
        objections: Optional[List[str]] = None
        lang: Optional[str] = "it"
        symptoms: List[str] = []
        kpis: List[str] = []
        triggers: Dict[str, List[str]] = {}
        raw_text: Optional[str] = None
        source: Optional[SourceRef] = None

# --- Richieste principali ----------------
class Contact(BaseModel):
    name: Optional[str] = None
    first_name: Optional[str] = None
    role: Optional[str] = None
    company: Optional[str] = None
    lang: Optional[str] = "it"
    email: Optional[EmailStr] = None
    company_id: Optional[str] = None

class ThreadSpec(BaseModel):
    name: str
    persona_id: Optional[str] = None
    channels: Optional[List[str]] = None
    cadence: Optional[List[Dict[str, Any]]] = None
    meta: Optional[Dict[str, Any]] = None

class Triggers(BaseModel):
    manual_priority: bool = False
    personal: Optional[str] = None
    competitor: Optional[List[str]] = None
    erp: Optional[List[str]] = None
    company_signal: Optional[str] = None
    linkedin_signal: Optional[str] = None
    read_inferred: Optional[str] = None

class ResearchOptions(BaseModel):
    mode: Literal['off', 'kb_only', 'web'] = 'web'  # default ON
    seeds: List[str] = Field(default_factory=list)
    max_sources: int = 5

class GenerateSequenceRequest(BaseModel):
    threads: Optional[List[ThreadSpec]] = None
    contacts: Optional[List[Contact]] = None
    buyer_persona_ids: Optional[List[str]] = None
    sequence_type: Optional[str] = "without_inmail"
    language: Optional[str] = "it"
    research: Optional[ResearchOptions] = None
    triggers: Optional[Triggers] = None
    force_variant: Optional[str] = None
    industry: Optional[str] = None
    role: Optional[str] = None

# --- Output composito / standard ---------

class WhatIUsed(BaseModel):
    personas: Optional[List[str]] = None
    files: Optional[List[str]] = None
    triggers: Optional[List[str]] = None

class StandardOutput(BaseModel):
    messages: Optional[List[Message]] = None
    coi: Optional[COI] = None
    sequence_next: Optional[List[SequenceAction]] = None
    calendar: Optional[List[CalendarEvent]] = None
    labels: Optional[List[str]] = None
    what_i_used: Optional[WhatIUsed] = None
    rationale: Optional[str] = None
    logging: Optional[dict] = None
    research: Optional[ResearchResult] = None

class RankItem(BaseModel):
    combo_id: str
    score: float
    rationale: Optional[str] = None

class RankRequest(BaseModel):
    """
    Richiesta di ranking: lista di candidati arbitrari (dict).
    """
    candidates: List[Dict[str, Any]] = Field(
        default_factory=list,
        description="Oggetti candidati (struttura libera) da ordinare/scorare."
    )

class RankResponse(BaseModel):
    """
    Risposta di ranking: lista ordinata di RankItem (score 0..1).
    """
    items: List[RankItem] = Field(default_factory=list)

#- --- Calendar build & ICS -------------------------------------------------------

class CalendarBuildResponse(BaseModel):
    """
    Ritorna la lista degli eventi e l'ICS (opzionale).
    """
    calendar: List[CalendarEvent] = Field(default_factory=list)
    ics: Optional[str] = Field(
        default=None,
        description="Contenuto file .ics (se richiesto/costruito)."
    )

class CalendarAction(BaseModel):
    day: int                      # offset dal "start_date"
    action: str                   # es. "email"|"linkedin_dm"|"inmail"
    channel: Optional[str] = None
    step: Optional[str] = None
    description: Optional[str] = None

class CalendarBuildRequest(BaseModel):
    start_date: Optional[date] = None
    rules: Optional[CalendarRules] = None
    base_sequence: List[CalendarAction] = Field(default_factory=list)

# --- COI stubs (se usi endpoints di estimate) ----------------------------------

try:
    HORIZON_DAYS_DEFAULT  # type: ignore[name-defined]
except NameError:
    HORIZON_DAYS_DEFAULT = 90

class ManualSignal(BaseModel):
    """
    Segnale manuale per la stima COI (nome/valore liberi).
    """
    name: str
    value: Optional[Any] = None

class COIEstimateRequest(BaseModel):
    """
    Input per stimare il COI su un certo orizzonte (giorni).
    """
    signals: Optional[List[ManualSignal]] = Field(default=None)
    horizon_days: conint = Field(default=HORIZON_DAYS_DEFAULT, ge=1, le=365)

class COIEstimateResponse(BaseModel):
    """
    Output stima COI.
    """
    status: Literal["ok", "error"] = "ok"
    estimate: Optional[float] = Field(
        default=None,
        description="Valore stimato (es. €) — se disponibile."
    )
    notes: Optional[List[str]] = Field(default=None)

# --- KB utility output

class KBSearchMatch(BaseModel):
    file: str
    score: confloat = Field(0.0, ge=0.0, le=1.0)
    snippet: Optional[str] = None

class KBSearchResponse(BaseModel):
    query: str
    matches: List[KBSearchMatch] = Field(default_factory=list)


class KBIngestRequest(BaseModel):
    url: Optional[str] = None
    content_base64: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None  # filename, tags, ecc.

class KBIngestResponse(BaseModel):
    doc_id: str
    chunks: int = 0
    urls: Optional[List[str]] = None
    text_preview: Optional[str] = None

class KBListItem(BaseModel):
    # superset per coprire gli usi diversi nel codice
    doc_id: Optional[str] = None
    file: Optional[str] = None
    title: Optional[str] = None
    lang: Optional[str] = None
    kind: Optional[str] = None
    size: Optional[int] = None
    metadata: Optional[Dict[str, Any]] = None

class KBListResponse(BaseModel):
    docs: List[KBListItem] = Field(default_factory=list)

class CompanyEvidence(BaseModel):
    company: Optional[str] = None
    industry: Optional[str] = None
    erp: Optional[List[str]] = None
    competitors: Optional[List[str]] = None
    signals: Optional[Dict[str, Any]] = None
    notes: Optional[str] = None

class UpsertResponse(BaseModel):
    ok: bool
    id: Optional[str] = None
    stored: Optional[Dict[str, Any]] = None

class Citation(BaseModel):
    url: str
    title: str
    snippet: Optional[str] = None
    published_at: Optional[datetime] = None
    source: Optional[str] = None

class EnrichedFact(BaseModel):
    text: str
    fact: Optional[str] = None
    confidence: float = 0.8
    source: Optional[str] = None
    citations: Optional[List[Citation]] = None

class ComposeStrategy(BaseModel):
    order: list[str] = Field(default_factory=list)
    apply_cleanup: bool = True

class CalendarRules(BaseModel):
    no_weekend: bool = True
    holiday_calendar: Optional[str] = None
    working_hours: Optional[dict] = None

class TriggerExtractResp(BaseModel):
    """
    Risultato estrazione trigger.
    - triggers: dizionario per categoria → lista stringhe
    - ranked: lista di combinazioni già valutate (score 0..1)
    """
    triggers: Dict[str, List[str]] = Field(default_factory=dict)
    ranked: List[Dict[str, Any]] = Field(default_factory=list)

class CompetitorRequest(BaseModel):
    company: str
    competitors: List[str]
    industry: Optional[str] = None
    icp: Optional[str] = None
    language: Optional[str] = "it"
    research: Optional[ResearchOptions] = None

class CompetitorResponse(BaseModel):
    company: str
    insights: List[str] = Field(default_factory=list)
    sources: List["SourceRef"] = Field(default_factory=list)
    matrix: "CompetitorMatrix"

class ExportSequenceResponse(BaseModel):
    ok: bool
    format: Literal["json", "csv", "docx"] = "json"
    sequence_id: Optional[str] = None
    path: Optional[str] = None


class CompetitorItem(BaseModel):
    name: str
    gtm: List[str] = []
    strengths: List[str] = []
    weaknesses: List[str] = []
    feature_coverage: Dict[str, float] = {}
    notes: List[str] = []
    sources: List[SourceRef] = []
    # scoring “arricchito”
    score_total: float = 0.0
    score_feature: float = 0.0
    score_industry_icp: float = 0.0

class CompetitorMatrix(BaseModel):
    company: str
    target: str
    industry: Optional[str] = None
    icp: Optional[str] = None
    normalized_features: List[str] = []
    feature_coverage: Dict[str, Dict[str, float]] = {}  # feature -> {competitor: coverage}
    items: List[CompetitorItem] = []
    research_summary: Optional[str] = None
    scoring: Dict[str, Any] = {}

HEADER_ALIASES = {
    "pain": ["pains", "pain points", "challenge", "challenges"],
    "kpi": ["kpi", "kpis", "metrics", "key metrics"],
    "symptom": ["symptoms", "signals", "triggers", "signs"],
    "objection": ["objections", "resistance", "concerns"],
}

class ResearchProviderError(Exception):
    pass

class SendEmailRequest(BaseModel):
    to: EmailStr
    subject: str
    text: str

class SendEmailResponse(BaseModel):
    ok: bool
    id: Optional[str] = None

class ExportDocxRequest(BaseModel):
    title: Optional[str] = None
    messages: List[Message]

class DiffRequest(BaseModel):
    a: List[Message]
    b: List[Message]

class DiffChunk(BaseModel):
    step: str
    channel: str
    diff_unified: str

class DiffResponse(BaseModel):
    chunks: List[DiffChunk]

class PersonaGenReq(BaseModel):
    industry: str
    role: str
    company_size: Optional[str] = None
    language: str = "it"

class PersonaGenResp(BaseModel):
    name: str
    pains: List[str]
    kpis: List[str]
    triggers: List[str]
    objections: List[str]
    messaging_notes: List[str]
    sources: List[SourceRef]

class PersonaRequest(BaseModel):
    industry: str
    role: str
    level: Optional[str] = None
    lang: Optional[str] = "it"

class DailyReminders(BaseModel):
    date: str
    items: List[Dict[str, Any]]

class DmRole(BaseModel):
    role: str
    cares_about: List[str]
    kpis: List[str]
    angle: Optional[str] = None

class DmMapResponse(BaseModel):
    industry: Optional[str] = None
    product: Optional[str] = None
    roles: List[DmRole]

class ObjectionInput(BaseModel):
    product: str
    context: Optional[str] = None
    currency: Optional[str] = "€"
    language: Optional[str] = "it"

class ObjectionResponses(BaseModel):
    responses: List[str]

class PersonaIndexItem(BaseModel):
    persona_id: str
    industry: Optional[str] = None
    role: Optional[str] = None
    pains: int
    kpis: int
    objections: int

class PersonaIndexResponse(BaseModel):
    items: List[PersonaIndexItem]

class ABPromoteRequest(BaseModel):
    sequence_id: str
    variant_id: str

class ABPromoteResponse(BaseModel):
    ok: bool
    sequence_id: str
    variant_id: str

class ExportDocxRequest(BaseModel):
    title: Optional[str] = None
    messages: List[Message]

class LinkedInCreativesReq(BaseModel):
    theme: str
    language: Optional[str] = "it"
    coi_enabled: Optional[bool] = True
    creative_enabled: Optional[bool] = True
    sector: Optional[str] = None
    base_persona: Optional[dict] = None  # se non passi persona_id
    persona_id: Optional[str] = None     # opzionale, se hai un registry
    evidence: Optional[List[str]] = None

class LinkedInCreativesRes(BaseModel):
    A_tipps: str
    B_poke: str

# ======================================================================
# Helper minimi (nomi citati dai warning, versioni “no-op” sicure)
# ======================================================================

# === 6) HELPER / STUB FUNCS USATE DAGLI ENDPOINT ===============================

def _unused_next_step_str(i: int) -> str:   # (fix per "_next_step_str is not defined")
    return f"D+{i}"

# --- ADD HERE: logging helpers (app.py, vicino ai modelli/utility) ---
def _drop_nones(d: dict) -> dict:
    return {k: v for k, v in (d or {}).items() if v is not None}

def _json_canonical(obj: object) -> str:
    try:
        return json.dumps(obj, sort_keys=True, ensure_ascii=False, separators=(",", ":"), allow_nan=False)
    except Exception:
        return json.dumps(str(obj), ensure_ascii=False, separators=(",", ":"))

def _wm_hash(payload: Any) -> str:
    # hash "morbido" per logging/idempotency
    try:
        s = json.dumps(payload, sort_keys=True, ensure_ascii=False)
    except Exception:
        s = str(payload)
    return str(abs(hash(s)))[:12]

def idem_get(key: Optional[str]) -> Optional[Any]:
    # se hai un’implementazione altrove, importa quella; questo spegne i warning
    return None

KB_INDEX = Path("kb/index.json")

# ====== BEGIN: kb folder loader ======
def _kb_sidecar_text(p: Path) -> str:
    """
    Se esiste un sidecar .txt accanto al JSON della KB, usalo per lo snippet.
    """
    try:
        side = p.with_suffix(".txt")
        if side.exists():
            return side.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        pass
    return ""

def kb_index_load() -> list[dict]:
    kb_dir = Path("kb")
    items = []
    for p in kb_dir.rglob("*.json"):
        try:
            obj = json.loads(p.read_text(encoding="utf-8"))
            title   = obj.get("title") or p.stem
            url     = obj.get("url","")
            snippet = (obj.get("snippet") or "").strip()
            if not snippet:
                side = _kb_sidecar_text(p)   # <— QUI
                if side:
                    snippet = side[:500]
            items.append({"title": title, "url": url, "snippet": snippet})
        except Exception:
            continue
    # opzionale: altri formati (.md/.txt)…
    return items

def get_kb_items_flat() -> list[dict]:
    # prova _list_kb_flat se esiste
    if "_list_kb_flat" in globals() and callable(globals().get("_list_kb_flat")):
        return _list_kb_flat()
    # fallback su _kb_items_flat se esiste
    if "_kb_items_flat" in globals() and callable(globals().get("_kb_items_flat")):
        return _kb_items_flat()
    return []


def _list_kb_flat() -> List[dict]:
    # Usa i tuoi sidecar .txt come corpus
    items = []
    for p in Path("kb/raw").glob("**/*"):
        if p.is_file() and p.suffix.lower() in (".txt",".md",".docx",".pdf"):
            try:
                side = p.with_suffix(".txt")
                txt = side.read_text(encoding="utf-8", errors="ignore") if side.exists() else p.name
                items.append({"title": p.stem, "url": str(p), "snippet": txt[:400]})
            except Exception:
                pass
    return items
# ====== END: kb folder loader ======

def kb_index_save(idx):
    KB_INDEX.write_text(json.dumps(idx, ensure_ascii=False, indent=2), encoding="utf-8")

def kb_get(doc_id):
    idx = kb_index_load()
    meta = idx.get(doc_id)
    if not meta: return None
    p = KB_RAW / f"{doc_id}.txt"
    if not p.exists(): return None
    return { "doc_id": doc_id, **meta, "content": p.read_text(encoding="utf-8") }

HEADER_ALIASES = {
    "pain": ["pains", "pain points", "challenge", "challenges"],
    "kpi": ["kpi", "kpis", "metrics", "key metrics"],
    "symptom": ["symptoms", "signals", "triggers", "signs"],
    "objection": ["objections", "resistance", "concerns"],
}

def _norm_header(h: str) -> str:
    h0 = (h or "").strip().lower()
    for canon, aliases in HEADER_ALIASES.items():
        if h0 == canon or any(h0 == a for a in aliases):
            return canon
    return h0

def _docx_to_text(path: str) -> str:
    try:
        import docx2txt
        return docx2txt.process(path) or ""
    except Exception:
        return ""

def build_persona_corpus(root: str = PERSONA_DIR) -> list[PersonaDoc]:
    """
    Scansiona ./kb/personas/*.docx e costruisce PersonaDoc normalizzati.
    Header riconosciuti (normalizzati): pain, kpi, symptom, objection, triggers
    """
    docs: list[PersonaDoc] = []
    p = Path(root)
    if not p.exists():
        return docs

    for f in p.glob("*.docx"):
        txt = _docx_to_text(str(f))
        if not txt:
            continue

        # estrazione semplice dei blocchi per header “normali”
        blocks: dict[str, list[str]] = {"pain": [], "kpi": [], "symptom": [], "objection": [], "triggers": []}
        current = None
        for raw in txt.splitlines():
            line = raw.strip()
            if not line:
                continue
            # header se finisce con ":" o è tutto maiuscolo “titolo”
            if line.endswith(":") or (len(line) <= 26 and line.isascii() and line == line.title()):
                current = _norm_header(line.replace(":", ""))
                continue
            if current and current in blocks:
                blocks[current].append(line)

        # meta da nome file: es. "persona_cio_retail.docx"
        stem = f.stem.lower()
        role_guess = None
        industry_guess = None
        parts = re.split(r"[_\-]", stem)
        for t in parts:
            if t in {"cio","cfo","cto","head","director","manager","ops","operations","procurement"}:
                role_guess = (role_guess or t)
            if t in {"retail","cpg","automotive","utilities","fsi","manufacturing"}:
                industry_guess = (industry_guess or t)

        docs.append(PersonaDoc(
            persona_id=stem,
            industry=industry_guess,
            role=role_guess,
            pains=blocks["pain"],
            objections=blocks["objection"],
            lang="it",
            symptoms=blocks["symptom"],
            kpis=blocks["kpi"],
            triggers={"raw": blocks["triggers"]} if blocks["triggers"] else {},
            raw_text=txt,
            source=SourceRef(title=f.stem, url=None, snippet=None, published_at=None),
        ))
    return docs

def _load_persona_index(force: bool = False) -> list[PersonaDoc]:
    global _PERSONA_INDEX
    if force or (_PERSONA_INDEX is None):
        _PERSONA_INDEX = build_persona_corpus()
    return _PERSONA_INDEX or []

def _refresh_persona_index():
    global _PERSONA_INDEX
    _PERSONA_INDEX = None
    try:
        _ensure_persona_index_loaded()
    except Exception:
        pass

def _ensure_persona_index_loaded(force: bool = False):
    """
    Alias di compatibilità: carica/ricarica l'indice persona e ritorna la lista.
    """
    return _load_persona_index(force=force)

# ===================== END PATCH: _perform_research =====================

def _summarize_sources_llm(hits: List[Dict[str, Any]], lang: str = "it") -> str:
    """
    Riassunto “ibrido”:
    - default: deterministico (titoli + keyword top).
    - se OPENAI_API_KEY presente e LLM_SUMMARY_ENABLED attivo, prova un riassunto LLM.
    """
    hits = hits or []
    titles = [h.get("title") for h in hits if h.get("title")]
    snippets = [h.get("snippet") for h in hits if h.get("snippet")]

    # baseline deterministica: top parole chiave dai titoli+snippet
    text_all = " ".join([*titles, *snippets]).lower()
    tokens = re.findall(r"[a-z0-9\-]{3,}", text_all)
    stop = set(["the","and","for","con","per","una","della","delle","degli","dei","del","che","con","this",
                "with","from","sono","nel","alla","alla","gli","le","dei","sia","anche","su","api","edi"])
    freq = Counter([t for t in tokens if t not in stop])
    top_kw = ", ".join([w for w,_ in freq.most_common(10)])

    base = f"Fonti analizzate: {len(hits)}.\n"
    if titles:
        base += ("Titoli: " if lang.startswith("it") else "Titles: ") + "; ".join(titles[:5]) + ".\n"
    base += ("Parole chiave: " if lang.startswith("it") else "Keywords: ") + top_kw + "."

    # opzionale: LLM se abilitato
    if os.getenv("LLM_SUMMARY_ENABLED","false").lower() in {"1","true","yes"} and os.getenv("OPENAI_API_KEY"):
        try:
            # implementazione leggera con OpenAI Responses API (se disponibile nell’ambiente)
            # evita dipendenze hard se non presenti.
            import openai  # type: ignore
            openai.api_key = os.getenv("OPENAI_API_KEY")
            prompt = (
                "Riepiloga in 5-7 frasi i punti chiave emersi da questi risultati di ricerca.\n\n" +
                "\n\n".join([f"- {h.get('title','')}: {h.get('snippet','')}" for h in hits[:8]])
            )
            model = os.getenv("OPENAI_MODEL","gpt-4o-mini")
            rsp = openai.ChatCompletion.create(  # compat per vecchie SDK
                model=model,
                messages=[{"role":"system","content":"Sei un ricercatore conciso."},
                          {"role":"user","content":prompt}]
            )
            llm_text = (rsp.choices[0].message["content"] or "").strip()
            if llm_text:
                return llm_text
        except Exception:
            pass

    return base

# ===================== BEGIN PATCH: trigger extractor (regex + sinonimi gerarchici) =====================
TRIGGER_SYNONYMS_HIER = {
    "industry:retail":    [r"\bretail(er|ing)?\b", r"\bgdo\b", r"\bgrocery\b", r"\bfashion\b", r"\bcpg\b"],
    "industry:automotive":[r"\bautomotive\b", r"\boem\b", r"\btier\s?[12]\b"],
    "industry:utilities": [r"\butilities?\b", r"\benergy\b", r"\bgas\b", r"\belectric(it(y|à))?\b"],
    "industry:manufacturing":[r"\bmanufactur(ing|er|ers)\b", r"\bindustry 4\.0\b"],
    "industry:fsi":       [r"\b(banking|insurance|financial services?)\b", r"\bfinserv\b"],

    "role:cio":           [r"\bCIO\b", r"\bChief Information Officer\b"],
    "role:it director":   [r"\bIT\s?(Director|Head)\b", r"\bHead of IT\b"],
    "role:it manager":    [r"\bIT\s?Manager\b", r"\bICT\s?Manager\b"],
    "role:supply chain":  [r"\bSupply\s?Chain\b", r"\blogistics?\b"],

    "erp:sap":            [r"\bSAP\b", r"\bS/?4HANA\b", r"\bIDOC\b", r"\bBAPI\b"],
    "erp:oracle":         [r"\bOracle\b", r"\bE-?Business Suite\b", r"\bFusion\b"],
    "erp:d365":           [r"\bDynamics( 365| D365)\b", r"\bBusiness Central\b", r"\bNAV\b"],
    "erp:as400":          [r"\bAS/?400\b", r"\bIBM i\b", r"\b(i)Series\b"],

    "tech:edi":           [r"\bEDI\b", r"\bEDIFACT\b", r"\bX12\b", r"\bAS2\b", r"\bOFTP2?\b"],
    "tech:api":           [r"\bAPI(s)?\b", r"\bREST\b", r"\bwebhook(s)?\b", r"\bGraphQL\b"],
    "tech:mft":           [r"\bMFT\b", r"\bManaged File Transfer\b", r"\bSFTP\b"],

    "company:merger":     [r"\bmerger\b", r"\bfusione\b", r"\bM&A\b"],
    "company:acquisition":[r"\bacquisit(io|ion)e?\b", r"\bacquired\b"],
    "company:hiring":     [r"\bhiring\b", r"\bassunzion(i|e)\b", r"\bstanno cercando\b"],
    "personal:promoted":  [r"\bpromoted?\b", r"\bpromoss[oa]\b"],
}

_TRIGGER_REGEX = {k: [re.compile(p, re.I) for p in pats] for k, pats in TRIGGER_SYNONYMS_HIER.items()}

def _extract_triggers_from_text(text: str) -> list[str]:
    t = (text or "")
    out: list[str] = []
    for key, regs in _TRIGGER_REGEX.items():
        if any(r.search(t) for r in regs):
            out.append(key)
    # dedup preservando ordine
    seen = set(); uniq=[]
    for x in out:
        if x not in seen:
            uniq.append(x); seen.add(x)
    return uniq

# ====== BEGIN: LinkedIn creative A/B with COI ======
COI_ENABLED = os.getenv("COI_ENABLED","true").lower() in {"1","true","yes"}

# ====================== BEGIN: POLISH + LINKEDIN CREATIVE ======================
import os, re

LLM_POLISH_ENABLED = os.getenv("LLM_POLISH_ENABLED","true").lower() in {"1","true","yes"}

def _simple_polish(text: str, persona_creative: dict) -> str:
    """Fallback deterministic: lessico preferito, rimozione jargon, frasi + corte."""
    if not text:
        return text
    pc = persona_creative or {}
    lex = pc.get("lexicon", {})
    avoid = set((lex.get("avoid") or []))
    preferred = set((lex.get("preferred") or []))

    # 1) rimuovi jargon/termini da evitare
    t = text
    for bad in avoid:
        t = re.sub(rf"\b{re.escape(bad)}\b", "", t, flags=re.IGNORECASE)

    # 2) comprimi spazi/punteggiatura doppia
    t = re.sub(r"\s{2,}", " ", t).strip()
    t = re.sub(r"\s+([,.;:!?])", r"\1", t)

    # 3) spezza righe lunghe (LinkedIn friendly): ~16 parole max
    out_lines = []
    for raw_line in t.splitlines():
        words = raw_line.split()
        chunk = []
        for w in words:
            chunk.append(w)
            if len(chunk) >= 16:
                out_lines.append(" ".join(chunk))
                chunk = []
        if chunk:
            out_lines.append(" ".join(chunk))
    t = "\n".join(out_lines).strip()

    # 4) se abbiamo "preferred", enfatizza inserendo al posto di termini generici comuni
    # (semplice: se trovi "digitalizzazione", sostituisci con "integrazione" ecc. — lasciamo soft)
    common_bad_to_pref = {
        "digitalizzazione": next(iter(preferred)) if preferred else "integrazione",
        "trasformazione":   next(iter(preferred)) if preferred else "integrazione"
    }
    for bad, good in common_bad_to_pref.items():
        t = re.sub(rf"\b{re.escape(bad)}\b", good, t, flags=re.IGNORECASE)

    return t.strip()

def _polish_with_persona_style(text: str, persona_creative: dict, language: str = "it") -> str:
    """
    Se LLM_POLISH_ENABLED e hai _llm_complete disponibile → usa LLM con stile persona.
    Altrimenti usa _simple_polish.
    """
    if not text:
        return text
    if LLM_POLISH_ENABLED and callable(globals().get("_llm_complete", None)):
        style_guide = {
            "tone": persona_creative.get("tone"),
            "style_guidelines": persona_creative.get("style_guidelines"),
            "lexicon": persona_creative.get("lexicon"),
            "emotional_triggers": persona_creative.get("emotional_triggers"),
        }
        prompt = f"""Rendi il testo seguente più chiaro, concreto e senza gergo.
                  Lingua: {language}.
                  Applica queste linee guida JSON (se presenti) senza cambiarne il senso:
                  {style_guide}
                  Mantieni emoji e formattazione LinkedIn-friendly (righe brevi).
                  Testo:
                  <<<
                  {text}
                  >>>"""
        try:
            return _llm_complete(prompt, temperature=0.2)  # tuo helper
        except Exception:
            pass
    # fallback deterministic
    return _simple_polish(text, persona_creative)

def generate_linkedin_creatives(theme: str,
                                persona_creative: dict,
                                evidence: list[str] | None = None,
                                language: str = "it",
                                coi_enabled: bool = True) -> dict:
    """
    Produce due varianti: A (TIPPS) e B (Poke).
    """
    ev = " • ".join([e.strip() for e in (evidence or []) if e.strip()])
    hooks = (persona_creative or {}).get("hooks", {})
    status_quo = (hooks.get("status_quo") or ["Processi di integrazione ancora artigianali?"])[0]
    poke_q = ((persona_creative or {}).get("poke_questions") or ["Cosa rende accettabile un 2–4% di errori EDI?"])[0]
    coi_line = "🧮 Se resta così, quanto costa nei prossimi 90 giorni?" if coi_enabled else ""

    A = f"""[{theme}]
👉 Trigger: {status_quo}
🔎 Insight: molte frizioni non sono nell’ERP ma nel time-to-integration e nel monitoraggio.
✅ Prova: {ev or "esempi: +18% ASN accuracy; -35% lead time onboarding"}
👉 Se può interessare, ti mostro un modo semplice per scalare senza stravolgere l’esistente. {coi_line}"""

    B = f"""[{theme}]
❓ {poke_q}
Mini-dato: {ev or "alcuni clienti hanno ridotto i fallimenti IDoc dal 6% al 2%"}
Se ha senso, condivido i 3 check per capire se il problema è davvero l’ERP… o solo il flusso. {coi_line}"""

    if (persona_creative or {}).get("enabled"):
        A = _polish_with_persona_style(A, persona_creative, language=language)
        B = _polish_with_persona_style(B, persona_creative, language=language)

    return {"A_tipps": A.strip(), "B_poke": B.strip()}


# ====================== END: POLISH + LINKEDIN CREATIVE ======================

FORBIDDEN_PATTERNS = [
    "dump prompt", "full prompt", "export prompt", "show your prompt",
    "framework mapping", "print all templates", "mostrami il prompt"
]

def _guard_ip(input_text: str):
    t = (input_text or "").lower()
    if any(p in t for p in FORBIDDEN_PATTERNS):
        raise HTTPException(status_code=403, detail="Request not allowed.")

# dentro i tuoi handler principali (prima di generare):
# Place this at the start of your endpoint function, before processing the request
_search_counters = defaultdict(int)
_search_reset_at = time.time() + 86400  # reset daily

def _provider():
    return settings.RESEARCH_PROVIDER

def _guard_budget(provider: str):
    global _search_reset_at
    now = time()
    if now > _search_reset_at:
        _search_counters.clear()
        _search_reset_at = now + 86400
    limits = {
        "brave": int(os.getenv("RESEARCH_BUDGET_BRAVE", "2000")),
        "serper": int(os.getenv("RESEARCH_BUDGET_SERPER", "2500")),
        "serpapi": int(os.getenv("RESEARCH_BUDGET_SERPAPI", "250")),
        "tavily": int(os.getenv("RESEARCH_BUDGET_TAVILY", "1000")),
        "searxng": 10**9,  # di fatto illimitato self-host
    }
    if _search_counters[provider] >= limits.get(provider, 1000):
        raise ResearchProviderError(f"Budget exhausted for {provider}")
    _search_counters[provider] += 1

# --- TEXT NORMALIZATION (shared for TF-IDF, competitor, matching) ---
import unicodedata
_URL_RE = re.compile(r"https?://\S+")
_WS_RE  = re.compile(r"\s+")

def _normalize_text(s: str) -> str:
    """
    Normalizzazione robusta:
    - to lower
    - rimozione URL, punteggiatura non alfanumerica
    - deaccent (NFKD)
    - compattazione whitespace
    """
    s = s or ""
    s = _URL_RE.sub(" ", s)
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = s.lower()
    # tieni lettere/numeri e pochi separatori semantici
    s = re.sub(r"[^a-z0-9\s\-_/\.]", " ", s)
    s = _WS_RE.sub(" ", s).strip()
    return s


def _normalize(items: List[Dict[str, Any]]) -> "ResearchResult":
    """items: list di dict normalizzati con chiavi: title, url, snippet."""
    facts = []
    citations = []
    triggers = {}  # opzionale: puoi popolarlo altrove

    for it in items:
        title = it.get("title") or ""
        url = it.get("url") or it.get("link") or ""
        snippet = it.get("snippet") or it.get("description") or ""
        if not url:
            continue
        facts.append(EnrichedFact(text=snippet or title, source=url, score=None))
        citations.append(Citation(url=url, title=title))
    return ResearchResult(triggers=triggers, facts=facts, citations=citations)

IDEM_DIR = Path(os.getenv("IDEMPOTENCY_DIR", ".cache/idempotency"))
IDEM_DIR.mkdir(parents=True, exist_ok=True)
IDEM_TTL_SECONDS = int(os.getenv("IDEMPOTENCY_TTL_SECONDS", "86400"))  # default 24h
_idem_lock = threading.Lock()

def _idem_path(key: str) -> Path:
    h = hashlib.sha256((key or "").encode("utf-8")).hexdigest()
    return IDEM_DIR / f"{h}.json"

def idem_get(key: Optional[str]) -> Optional["StandardOutput"]:
    if not key:
        return None
    p = _idem_path(key)
    if not p.exists():
        return None
    try:
        with _idem_lock:
            data = json.loads(p.read_text(encoding="utf-8"))
        ts = float(data.get("_ts", 0))
        if (time.time() - ts) > IDEM_TTL_SECONDS:
            p.unlink(missing_ok=True)
            return None
        return StandardOutput.parse_obj(data["out"])
    except Exception:
        return None

def idem_set(key: Optional[str], out: "StandardOutput") -> None:
    if not key:
        return
    blob = {"_ts": time.time(), "out": out.dict()}
    p = _idem_path(key)
    with _idem_lock:
        p.write_text(json.dumps(blob, ensure_ascii=False), encoding="utf-8")

def _unused_next_step_str(i: int) -> str:
    return f"step-{i:02d}"

def _seed_from(text: str) -> int:
    import zlib
    return zlib.adler32((text or "").encode("utf-8"))

try:
    from docx import Document as _Docx
except Exception:
    _Docx = None


_AB_OPENERS = [
    "Buongiorno {first_name},",
    "Ciao {first_name},",
    "Gentile {first_name},",
]
_AB_CTAS = [
    "Se utile, posso condividere una check-list in 12’ — mercoledì le va?",
    "Le va un confronto rapido di 12’? Mercoledì o giovedì?",
    "Se ha senso, 10’ domani per una review rapida?",
]

_AB_CTAS_EMAIL = [
    "Le va un confronto rapido di 12’ mercoledì o giovedì?",
    "Se utile, condivido una check-list in 12’. Ha senso questa settimana?",
]
_AB_CTAS_LI = [
    "Se ha senso, 10’ domani per una review rapida?",
    "Posso condividere 3 check in 12’ — ha piacere?",
]

def _apply_openers_cta_safe(msg: "Message", c: Optional["Contact"], seed: int) -> "Message":
    try:
        return _apply_openers_cta(msg, c, seed)  # usa quella avanzata se esiste
    except NameError:
        opener = _AB_OPENERS[seed % len(_AB_OPENERS)]
        fn = (getattr(c, "first_name", "") or getattr(c, "name", "")).split(" ")[0] if c else ""
        opener = opener.format(first_name=(fn or "")).strip()

        ctas = _AB_CTAS_LI if msg.channel in ("linkedin_dm","inmail") else _AB_CTAS_EMAIL
        cta = ctas[(seed // 7) % len(ctas)]

        body = (msg.text or "").strip()
        if not body.lower().startswith(("ciao", "buongiorno", "gentile")):
            body = opener + "\n\n" + body
        if "?" not in body[-140:]:
            body = body + "\n\n" + cta
        try:
            return msg.model_copy(update={"text": body})
        except Exception:
            msg.text = body
            return msg

def _apply_openers_cta(m: "Message", contact: "Contact", seed: int) -> "Message":
    opener = _AB_OPENERS[seed % len(_AB_OPENERS)]
    cta    = _AB_CTAS[(seed // 7) % len(_AB_CTAS)]  # cambia “fase”
    fn     = (contact.first_name or contact.name or "").split(" ")[0]
    opener = opener.format(first_name=(fn or "")).strip()

    body = (m.text or "").strip()
    if not body.lower().startswith(("ciao", "buongiorno", "gentile")):
        body = opener + "\n\n" + body
    if "?" not in body[-140:]:
        body = body + "\n\n" + cta
    return m.model_copy(update={"text": body})

import re

def export_sequence_to_docx(messages: List["Message"], title: str = "FlowAgent v3 — Sequenza") -> str:
    if _Docx is None:
        raise HTTPException(501, "python-docx non disponibile")
    doc = _Docx()
    doc.add_heading(title, level=1)
    for m in messages or []:
        doc.add_heading(f"{m.channel} · {m.step}", level=2)
        if m.subject:
            doc.add_paragraph(f"Oggetto: {m.subject}")
        doc.add_paragraph(m.text or "")
        if getattr(m, "tips", None):
            doc.add_paragraph("Tips: " + ", ".join(m.tips))

    out_path = Path("exports")
    out_path.mkdir(parents=True, exist_ok=True)

    safe_title = re.sub(r"[^\w-]+", "_", (title or "untitled").lower())
    file_path = out_path / f"{safe_title}.docx"

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    file_path.write_bytes(bio.read())
    return str(file_path)


def full_gpt_instructions_path() -> str | None:
    path = os.path.join("kb", "istruzioni_gpt_full.pdf")
    return path if os.path.exists(path) else None

@app.middleware("http")
async def ip_whitelist(request: Request, call_next):
    _guard_ip_from_request(request)
    return await call_next(request)

# ===================== BEGIN PATCH: persona_lookup_advanced =====================
# (richiede che _load_persona_index() e PersonaDoc esistano)
def persona_lookup_advanced(triggers: List[str], top_k: int = 3) -> List[Dict[str, Any]]:
    """
    Ranking semplice delle buyer personas presenti a KB in base ai trigger.
    Scoring: +1 per match industry/role, +0.5 per keyword in pains/kpis/symptoms, +0.25 tech/erp.
    """
    idx = _load_persona_index() or []
    if not triggers:
        # fallback: ordina per presenza di industry/role (qualunque)
        return [{"persona_id": p.persona_id, "score": 0.0} for p in idx][:top_k]

    trig_set = set([t.lower() for t in triggers])
    scored: List[Tuple[str, float]] = []
    for p in idx:
        s = 0.0
        # industry/role
        if p.industry:
            if f"industry:{p.industry}".lower() in trig_set: s += 1.0
        if p.role:
            if f"role:{p.role}".lower() in trig_set: s += 1.0
        # ERP/tech
        for tag in ("sap","oracle","nav","bc","d365","edi","api","webhook","graphql","idoc"):
            if any(tag in t for t in trig_set):
                s += 0.25
                break
        # pains/kpis/symptoms keyword overlap
        bag = " ".join([*(p.pains or []), *(p.kpis or []), *(p.symptoms or [])]).lower()
        for t in trig_set:
            val = t.split(":",1)[1] if ":" in t else t
            if len(val) >= 3 and val in bag:
                s += 0.5
        scored.append((p.persona_id, s))

    scored.sort(key=lambda x: x[1], reverse=True)
    return [{"persona_id": pid, "score": float(sc)} for pid, sc in scored[:top_k]]
# ===================== END PATCH: persona_lookup_advanced =====================


# ===================== BEGIN PATCH: pick_best_combo =====================
# (aggiungi in testa se manca)
# from typing import Optional

# mappa minima trigger→KPI (estendila nel tempo)

PERSONA_KPI_MAP = {
    # pesi indicativi: adatta nel tempo
    "role:cio": {
        "erp:sap": {"integration-lead-time": 0.6, "idoc-success-rate": 0.4},
        "tech:edi": {"onboarding-lead-time": 0.5, "as2-success-rate": 0.5},
        "*": {"integration-cost": 0.7, "system-availability": 0.3},
    },
    "role:it director": {
        "tech:api": {"time-to-first-integration": 0.7, "defect-leakage": 0.3},
        "*": {"backlog-burn-down": 0.6, "change-failure-rate": 0.4},
    },
    "industry:retail": {
        "tech:edi": {"on-time-fulfillment": 0.6, "return-rate": 0.4},
        "*": {"ticket-rate": 0.5, "stockout-rate": 0.5},
    },
    "industry:automotive": {
        "tech:edi": {"asn-accuracy": 0.6, "ppap-cycle-time": 0.4},
        "*": {"supplier-onboarding-time": 0.5, "delivery-otd": 0.5},
    },
}

def _select_kpis(triggers: list[str], persona_id: str | None) -> list[tuple[str,float]]:
    trig_set = set(triggers or [])
    buckets: dict[str,float] = {}
    # 1) persona rule (role/industry) se presente nel trigger set
    for key, sub in PERSONA_KPI_MAP.items():
        if key in trig_set:
            # try matching by more specific combo (e.g. erp:/tech:)
            matched = False
            for t in trig_set:
                if t.startswith("erp:") or t.startswith("tech:"):
                    if t in sub:
                        for kpi, w in sub[t].items():
                            buckets[kpi] = buckets.get(kpi, 0.0) + w
                        matched = True
            if not matched and "*" in sub:
                for kpi, w in sub["*"].items():
                    buckets[kpi] = buckets.get(kpi, 0.0) + w
    # 2) fallback assoluto
    if not buckets:
        for kpi, w in {"integration-lead-time":0.5,"integration-cost":0.5}.items():
            buckets[kpi] = buckets.get(kpi, 0.0) + w
    # normalize
    tot = sum(buckets.values()) or 1.0
    return sorted([(k, round(v/tot,3)) for k,v in buckets.items()], key=lambda x: x[1], reverse=True)

def pick_best_combo(triggers: list[str],
                    persona_fit: list[dict[str, any]],
                    hits: list[dict[str, any]] | None = None) -> dict[str, any]:
    persona_id = (persona_fit[0]["persona_id"] if persona_fit else "generic")
    # priorità: industry > erp > tech > altro
    pref = [t for t in triggers if t.startswith("industry:")] or \
           [t for t in triggers if t.startswith("erp:")] or \
           [t for t in triggers if t.startswith("tech:")] or \
           (triggers[:1] if triggers else [])
    chosen = pref[0] if pref else None
    # KPI pesati
    kpis_w = _select_kpis(triggers, persona_id)
    return {
        "persona_id": persona_id,
        "trigger": chosen,
        "kpis": [k for k, _ in kpis_w],
        "kpi_weights": {k:w for k,w in kpis_w},
        "evidence_count": len(hits or []),
    }
# ===================== END PATCH: pick_best_combo =====================

# --- Utils comuni (dedup/normalize) ---

def _ensure_cache_dirs():
    from pathlib import Path
    (Path(".cache/kb_query")).mkdir(parents=True, exist_ok=True)
    (Path(".cache/url_text")).mkdir(parents=True, exist_ok=True)
_ensure_cache_dirs()


# ===================== BEGIN PATCH: URL->text cache helper =====================
# --- KB query cache (unica versione) ---
CACHE_DIR = Path(os.getenv("CACHE_DIR", ".cache"))
KB_QCACHE_DIR = CACHE_DIR / "kb_query"
KB_QCACHE_DIR.mkdir(parents=True, exist_ok=True)

_KB_QCACHE_FILE = KB_QCACHE_DIR / "qcache.json"   # nome coerente con le funzioni legacy se le tieni
KB_QCACHE_TTL = int(os.getenv("KB_QCACHE_TTL", "86400"))  # 24h default



URL_TXT_DIR = CACHE_DIR / "url_text"
URL_TXT_DIR.mkdir(parents=True, exist_ok=True)
URL_TXT_TTL = int(os.getenv("URL_TXT_TTL", "2592000"))  # 30 giorni

# app.py
def _qcache_path(key: dict) -> Path:
    h = hashlib.sha256(json.dumps(key, sort_keys=True, ensure_ascii=False).encode("utf-8")).hexdigest()
    return KB_QCACHE_DIR / f"{h}.json"

def research_cache_get(key: dict):
    p = _qcache_path(key)
    if not p.exists(): return None
    try:
        obj = json.loads(p.read_text(encoding="utf-8"))
        # opzionale TTL:
        return obj.get("items", None)
    except Exception:
        return None

def research_cache_set(key: dict, items: list[dict]):
    p = _qcache_path(key)
    try:
        p.write_text(json.dumps({"items": items}, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass

def _urltxt_path(url: str) -> Path:
    h = hashlib.sha256((url or "").encode("utf-8")).hexdigest()
    return URL_TXT_DIR / f"{h}.json"

def url_text_cache_get(url: str) -> str | None:
    p = _urltxt_path(url)
    if not p.exists(): return None
    try:
        obj = json.loads(p.read_text(encoding="utf-8"))
        if (time.time() - float(obj.get("_ts", 0))) > URL_TXT_TTL:
            p.unlink(missing_ok=True); return None
        return obj.get("text") or None
    except Exception:
        return None

def url_text_cache_put(url: str, text: str):
    p = _urltxt_path(url)
    try:
        p.write_text(json.dumps({"_ts": time.time(), "text": text or ""}, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass

_KB_QCACHE_FILE = Path(".cache/kb_query/qcache.json")
try:
    if not _KB_QCACHE_FILE.exists():
        _KB_QCACHE_FILE.write_text("{}", encoding="utf-8")
except Exception:
    pass

def _kb_items_flat() -> List[dict]:
    """Corpus KB: usa i sidecar .txt accanto ai file in kb/raw."""
    items: List[dict] = []
    for p in Path("kb/raw").glob("**/*"):
        if not p.is_file():
            continue
        if p.suffix.lower() not in (".txt", ".md", ".docx", ".pdf"):
            continue
        try:
            side = p.with_suffix(".txt")
            txt = side.read_text(encoding="utf-8", errors="ignore") if side.exists() else p.name
            items.append({"title": p.stem, "url": str(p), "snippet": (txt or "")[:500]})
        except Exception:
            pass
    return items


# ===================== END PATCH: URL->text cache helper =====================


def _has_action(actions: List["SequenceAction"], step: str, channel: str) -> bool:
    return any((a.step == step and a.channel == channel) for a in actions)

def plan_inmail_bump_if_no_connection(
    messages: List["Message"],
    actions: List["SequenceAction"],
    calendar: List["CalendarEvent"],
    thread: ThreadSpec,
    payload: Any,
    start_date: date,
    n_days_after_dm: int = 3,
    li_status_resolver=None,
):
    _ = thread, payload, li_status_resolver  # silence not-used
    """
    Se c'è un DM LinkedIn (linkedin_dm) allo step N, pianifica una InMail bump allo step N+1
    dopo n_days_after_dm giorni. Non duplica se già presente.
    (Status resolver non usato per ora → sempre pianifica la bump deterministica.)
    """
    # trova l’ultimo step DM
    dm_steps = sorted({m.step for m in messages if m.channel == "linkedin_dm"})
    if not dm_steps:
        return
    last_dm_step = dm_steps[-1]
    # step successivo
    try:
        idx = int(last_dm_step.replace("step-",""))
        next_step = f"step-{idx+1:02d}"
    except Exception:
        next_step = f"{last_dm_step}-bump"

    if _has_action(actions, next_step, "inmail"):
        return

    actions.append(SequenceAction(channel="inmail", step=next_step, action="send"))
    calendar.append(CalendarEvent(date=start_date + timedelta(days=max(1, n_days_after_dm)),
                                  action="inmail_bump"))

def plan_dm_followup_on_connection_accept(
    messages: List["Message"],
    actions: List["SequenceAction"],
    calendar: List["CalendarEvent"],
    thread: "ThreadSpec",
    payload: Any,
    start_date: date,
    n_days_after_dm: int = 3,
    li_status_resolver=None,
):
    _ = thread, payload, li_status_resolver  # silence not-used

    """
    Pianifica un follow-up DM dopo l’ultimo DM, n_days_after_dm giorni dopo.
    (Per ora non consulta lo stato reale della connessione.)
    """
    dm_steps = sorted({m.step for m in messages if m.channel == "linkedin_dm"})
    if not dm_steps:
        return
    last_dm_step = dm_steps[-1]
    try:
        idx = int(last_dm_step.replace("step-",""))
        next_step = f"step-{idx+1:02d}"
    except Exception:
        next_step = f"{last_dm_step}-followup"

    if _has_action(actions, next_step, "linkedin_dm"):
        return

    actions.append(SequenceAction(channel="linkedin_dm", step=next_step, action="send"))
    calendar.append(CalendarEvent(date=start_date + timedelta(days=max(1, n_days_after_dm)),
                                  action="dm_followup"))

# ===================== BEGIN PATCH: _polish_with_llm =====================
def _polish_with_llm(messages: List["Message"], lang: str = "it") -> List["Message"]:
    """
    Fase di “polish”:
    - baseline deterministica: normalizza soggetto, aggiunge una CTA chiara se assente.
    - se OPENAI_API_KEY presente e flag abilitato → prova a riscrivere i testi con tono conciso.
    """
    def _deterministic_tweak(m: "Message") -> "Message":
        subj = m.subject
        if subj:
            subj = subj.strip().capitalize()
            if not subj.endswith("?") and "?" in subj:
                pass
        body = (m.text or "").strip()
        # CTA se manca
        if "12" not in body and "15" not in body and "min" not in body and " minuti" not in body:
            body += ("\n\n" + ("Ha senso un confronto di 12’ questa settimana?" if lang.startswith("it")
                               else "Would a 12’ chat this week make sense?"))
        return type(m)(**{**m.__dict__, "subject": subj, "text": body})

    polished = [_deterministic_tweak(m) for m in messages]

    if os.getenv("LLM_POLISH_ENABLED","false").lower() in {"1","true","yes"} and os.getenv("OPENAI_API_KEY"):
        try:
            import openai  # type: ignore
            openai.api_key = os.getenv("OPENAI_API_KEY")
            model = os.getenv("OPENAI_MODEL","gpt-4o-mini")
            out = []
            for m in polished:
                prompt = (
                    "Rendi il seguente messaggio conciso, chiaro e orientato all’azione. Mantieni il senso.\n\n"
                    f"Soggetto: {m.subject or '(nessuno)'}\n\nTesto:\n{m.text}\n"
                )
                rsp = openai.ChatCompletion.create(
                    model=model,
                    messages=[{"role":"system","content":"Sei un assistente di sales enablement."},
                              {"role":"user","content":prompt}]
                )
                new_txt = (rsp.choices[0].message["content"] or "").strip()
                # mantieni subject se non è stato suggerito un soggetto
                out.append(type(m)(**{**m.__dict__, "text": new_txt or m.text}))
            polished = out
        except Exception:
            pass
    return polished
# ===================== END PATCH: _polish_with_llm =====================


def _default_cadence_for(sequence_type: str | None) -> list[dict]:
    """
    Ritorna una cadence di default in assenza di thread.cadence.
    sequence_type: "with_inmail", "without_inmail" o None
    """
    if (sequence_type or "").lower() == "with_inmail":
        return [
            {"day": 0, "action": "email"},
            {"day": 3, "action": "linkedin_dm"},
            {"day": 7, "action": "inmail"},
        ]
    else:
        return [
            {"day": 0, "action": "email"},
            {"day": 3, "action": "linkedin_dm"},
            {"day": 7, "action": "email_followup"},
        ]

def _normalize_name(s: str) -> str:
    import re as _re
    return _re.sub(r"[^a-z0-9]+", "", (s or "").lower())

# ================== BEGIN: get_persona_by_id ==================
def get_persona_by_id(persona_id: str, log: dict | None = None) -> Optional[dict]:
    """
    Cerca una persona per ID (stem del file .docx).
    Fallback: se non trovata, prova a ricostruire da industry/role noti.
    """
    if not persona_id:
        return None
    pid = (persona_id or "").strip().lower()
    for p in _load_persona_index() or []:
        if (p.persona_id or "").strip().lower() == pid:
            # ritorna in forma dict per compatibilità con gli endpoint content
            return {
                "persona_id": p.persona_id,
                "industry": p.industry,
                "role": p.role,
                "pains": p.pains,
                "kpis": p.kpis,
                "objections": p.objections,
                "symptoms": p.symptoms,
                "triggers": p.triggers,
                "raw_text": p.raw_text,
            }
    # Fallback “closest” se l’ID è tipo "cio_automotive"
    try:
        parts = re.split(r"[_\-]", pid)
        ind = next((t for t in parts if t in {"retail","cpg","automotive","utilities","fsi","manufacturing"}), None)
        rol = next((t for t in parts if t in {"cio","cfo","cto","director","manager","ops"}), None)
        best_pid, why, sc = _closest_persona_id(ind, rol)
        if best_pid:
            if isinstance(log, dict):
                log.setdefault("persona_fallback_debug", {})["closest"] = {"why": why, "score": sc}

            return get_persona_by_id(best_pid)
    except Exception:
        pass
    return None
# ================== END: get_persona_by_id ==================


def _closest_persona_id(industry: Optional[str], role: Optional[str]) -> Tuple[Optional[str], str, float]:
    # pesi: 60% industry, 40% role (configurabili)
    w_ind = float(os.getenv("PERSONA_MATCH_W_INDUSTRY", "0.6"))
    w_role = float(os.getenv("PERSONA_MATCH_W_ROLE", "0.4"))
    s = w_ind + w_role
    if s <= 0: w_ind, w_role = 0.6, 0.4
    else: w_ind, w_role = (w_ind/s), (w_role/s)

    best_pid, best_why, best_score = None, "no_match", 0.0
    for pid, ind, rol in _list_known_personas():
        score = w_ind * _string_sim(industry, ind) + w_role * _string_sim(role, rol)
        if score > best_score:
            best_pid, best_score = pid, score
            best_why = f"closest:{ind or 'n/a'}/{rol or 'n/a'}"
    return best_pid, best_why, best_score

DEFAULT_FEATURES: list[str] = [
    "edi", "as2", "api", "mft", "mapping", "translator", "portal", "onboarding",
    "monitoring", "alerts", "sla", "retry", "sftp", "security", "compliance", "sap", "oracle", "netsuite"
]
_FEATURE_SYNONYMS = {
    "edi": [
        r"\bedi\b", r"\b(electronic\s+data\s+interchange)\b",
        r"\btrading\s*partner\b", r"\bas2\b", r"\b(x12|edifact)\b"
    ],
    "api": [
        r"\bapi(s)?\b", r"\brest(-?api)?\b", r"\bgraphql\b", r"\bwebhook(s)?\b"
    ],
    "erp": [
        r"\berp\b", r"\bsap\b", r"\bs/4hana\b", r"\boracle\s+e-?business\b", r"\bmicrosoft\s+dynamics\b", r"\bnav(ision)?\b"
    ],
    "onboarding": [
        r"\bonboarding\b", r"\bpartner\s+onboarding\b", r"\bsupplier\s+enable(ment|ment)\b", r"\bvendor\s+onboarding\b"
    ],
    "lead_time": [
        r"\blead[-\s]?time\b", r"\btime[-\s]?to[-\s]?value\b", r"\btime[-\s]?to[-\s]?market\b"
    ],
    "compliance": [
        r"\bcompliance\b", r"\bgx?p\b", r"\baudit(s)?\b", r"\biso\s?\d{3,}\b", r"\bsoc\s?2\b"
    ],
    "security": [
        r"\bsecurity\b", r"\bcyber(security)?\b", r"\bencryption\b", r"\bpki\b", r"\btls\b"
    ],
    "integration": [
        r"\bintegration(s)?\b", r"\bmiddleware\b", r"\besb\b", r"\bipaa?s\b", r"\bipaas\b", r"\b(etl|elt)\b"
    ],
    "supply_chain": [
        r"\bsupply\s+chain\b", r"\blogistic(s)?\b", r"\b3pl\b", r"\bwarehouse\s+mgmt\b", r"\bwms\b", r"\btms\b"
    ],
    "forecasting": [
        r"\bforecast(ing)?\b", r"\bdemand\s+plan(ning)?\b", r"\bs&op\b"
    ],
    "crm": [
        r"\bcrm\b", r"\bsalesforce\b", r"\bhubspot\b", r"\bdynamics\s+365\s+sales\b"
    ],
    "omnichannel": [
        r"\bomni[-\s]?channel\b", r"\bclick[-\s]?and[-\s]?collect\b", r"\bship[-\s]?from[-\s]?store\b"
    ],
}

def _feature_hit_score(feature: str, text: str) -> float:
    pats = _FEATURE_SYNONYMS.get((feature or "").lower(), [])
    if not pats:
        return 0.0
    hits = 0
    for p in pats:
        if re.search(p, text, flags=re.I):
            hits += 1
    # normalizza su numero di pattern per feature
    return min(1.0, hits / max(1, len(pats)))

def _normalize(items: List[Dict[str, Any]]) -> "ResearchResult":
    facts, citations = [], []
    for it in items:
        title = it.get("title") or ""
        url = it.get("url") or it.get("link") or ""
        snippet = it.get("snippet") or it.get("description") or ""
        if not url:
            continue
        facts.append(EnrichedFact(text=(snippet or title), source=url, confidence=0.8))
        citations.append(Citation(url=url, title=title))
    return ResearchResult(triggers={}, facts=facts, citations=citations)

def _score_feature_coverage(feature: str, corpus: str) -> float:
    f = (feature or "").lower().strip()
    t = (corpus or "").lower()
    # sinonimi base
    SYNS: Dict[str, List[str]] = {
        "edi": ["edi", "electronic data interchange", "x12", "edifact"],
        "api": ["api", "rest", "graphql", "webhook", "endpoint"],
        "mft": ["mft", "managed file transfer", "sftp", "ftps"],
        "mapping": ["mapping", "mapper", "trasform", "translator", "idoc map"],
        "onboarding": ["onboarding", "partner onboarding", "trading partner", "supplier onboarding"],
        "monitoring": ["monitor", "monitoring", "observability", "alert"],
    }
    patterns = SYNS.get(f, [f])
    hits = sum(t.count(p) for p in patterns)
    if hits == 0: return 0.0
    if hits == 1: return 0.6
    return 1.0

def _score_industry_icp(corpus: str, industry: Optional[str], icp: Optional[str]) -> float:
    if not corpus or (not industry and not icp):
        return 0.5
    score = 0.5
    if industry and re.search(re.escape(industry), corpus, re.I):
        score += 0.25
    if icp and re.search(re.escape(icp), corpus, re.I):
        score += 0.25
    return float(max(0.0, min(1.0, score)))

def _build_corpus_per_competitor(research_hits: list[dict], competitors: list[str]) -> dict[str, str]:
    """
    Crea un corpus testuale per competitor concatenando titoli/snippet/url (deduplicati).
    """
    corpora: dict[str, list[str]] = {c.lower(): [] for c in (competitors or [])}
    seen = set()
    for h in (research_hits or []):
        title = _normalize_text(h.get("title"))
        snippet = _normalize_text(h.get("snippet"))
        url = _normalize_text(h.get("url"))
        blob = " ".join([title, snippet, url]).strip()
        if not blob:
            continue
        key = (title, url)
        if key in seen:
            continue
        seen.add(key)
        for c in corpora.keys():
            # match semplice: il nome del competitor appare nel titolo/snippet/url
            if c and (c in title or c in snippet or c in url):
                corpora[c].append(blob)
    return {c: " ".join(parts) for c, parts in corpora.items()}

def _score_competitor_item(
    item: "CompetitorItem",
    normalized_features: list[str],
    corpus: str,
    industry: Optional[str],
    icp: Optional[str],
    weight_feature: float = 0.7,
    weight_industry_icp: float = 0.3
) -> "CompetitorItem":
    cov = []
    for f in normalized_features:
        val = item.feature_coverage.get(f, None)
        if val is None:
            val = _score_feature_coverage(f, corpus)
            item.feature_coverage[f] = val
        cov.append(val)
    score_feature = sum(cov) / max(1, len(cov))
    score_ind_icp = _score_industry_icp(corpus, industry, icp)
    total = weight_feature * score_feature + weight_industry_icp * score_ind_icp

    item.score_feature = round(score_feature, 3)
    item.score_industry_icp = round(score_ind_icp, 3)
    item.score_total = round(total, 3)
    return item

def _split_corpus_by_competitor(sources: List[Dict[str, Any]], competitors: List[str]) -> Dict[str, str]:
    m = {c.lower(): [] for c in (competitors or [])}
    for h in (sources or []):
        text = " ".join([h.get("title") or "", h.get("snippet") or ""])
        for comp in m.keys():
            if comp in (text.lower()):
                m[comp].append(text)
    return {k: "\n".join(v) for k, v in m.items()}

# Feature “vocabolario” normalizzato (puoi arricchire)
DEFAULT_FEATURES = [
    "as2", "sftp", "mft", "api", "mapping", "partner_onboarding",
    "self_service", "dashboard", "monitoring", "sla", "pricing", "templates"
]
# Pattern di copertura (regex) per ogni feature
FEATURE_PATTERNS: Dict[str, List[re.Pattern]] = {
    "as2": [re.compile(r"\bAS2\b", re.I)],
    "sftp": [re.compile(r"\bSFTP\b", re.I)],
    "mft": [re.compile(r"\bMFT\b", re.I), re.compile(r"managed file transfer", re.I)],
    "api": [re.compile(r"\bAPI\b", re.I), re.compile(r"\bREST\b", re.I), re.compile(r"\bwebhook(s)?\b", re.I)],
    "mapping": [re.compile(r"\bmapping(s)?\b", re.I), re.compile(r"\btranslator(s)?\b", re.I)],
    "partner_onboarding": [re.compile(r"\bonboard(ing)? partner(s)?\b", re.I),
                           re.compile(r"\btrading partner(s)?\b", re.I)],
    "self_service": [re.compile(r"\bself[- ]service\b", re.I), re.compile(r"\bportal\b", re.I)],
    "dashboard": [re.compile(r"\bdashboard(s)?\b", re.I)],
    "monitoring": [re.compile(r"\bmonitor(ing)?\b", re.I), re.compile(r"\balert(s|ing)?\b", re.I)],
    "sla": [re.compile(r"\bSLA(s)?\b", re.I), re.compile(r"\bservice level\b", re.I)],
    "pricing": [re.compile(r"\bpricing\b", re.I), re.compile(r"\bprice\b", re.I)],
    "templates": [re.compile(r"\btemplate(s)?\b", re.I), re.compile(r"\bcatalog(ue)?\b", re.I)],
}

FEATURE_PATTERNS.update({
    "as2": FEATURE_PATTERNS.get("as2", []) + [re.compile(r"\bApplicability Statement 2\b", re.I)],
    "sftp": FEATURE_PATTERNS.get("sftp", []) + [re.compile(r"\bSSH File Transfer Protocol\b", re.I)],
    "api": FEATURE_PATTERNS.get("api", []) + [re.compile(r"\bOpenAPI\b", re.I), re.compile(r"\bSwagger\b", re.I)],
    "mapping": FEATURE_PATTERNS.get("mapping", []) + [re.compile(r"\btranslat(or|ion)\b", re.I)],
    "partner_onboarding": FEATURE_PATTERNS.get("partner_onboarding", []) + [re.compile(r"\bself-?service onboarding\b", re.I)],
    "dashboard": FEATURE_PATTERNS.get("dashboard", []) + [re.compile(r"\bcontrol tower\b", re.I)],
    "monitoring": FEATURE_PATTERNS.get("monitoring", []) + [re.compile(r"\balert(ing)?\b", re.I)],
    "sla": FEATURE_PATTERNS.get("sla", []) + [re.compile(r"\bservice level(s)?\b", re.I)],
    "pricing": FEATURE_PATTERNS.get("pricing", []) + [re.compile(r"\blicensing\b", re.I)],
    "templates": FEATURE_PATTERNS.get("templates", []) + [re.compile(r"\baccelerators?\b", re.I)],
})

def _norm_features(payload) -> list:
    """
    Estrae la lista delle feature normalizzate dal payload CompetitorRequest.
    Usa DEFAULT_FEATURES come fallback se non specificato.
    """
    if hasattr(payload, 'normalized_features') and payload.normalized_features:
        return list(payload.normalized_features)
    if hasattr(payload, 'features') and payload.features:
        return list(payload.features)
    return list(DEFAULT_FEATURES)



async def build_competitor_matrix(payload: "CompetitorRequest") -> "CompetitorMatrix":
    """
    Coverage matrix deterministica (KB/Web in base a payload.research.mode) con:
      - normalized_features
      - feature_coverage (+ bonus regex)
      - ranking (pesi industry/ICP)
      - research_summary (se LLM_SUMMARY_ENABLED)
    """

    # 0) dedup nomi competitor

    uniq, seen = [], set()
    for comp in (payload.competitors or []):
        k = _normalize_name(comp)
        if k in seen:
            continue
        seen.add(k); uniq.append(comp)
    payload.competitors = uniq

    # …dentro generate_sequence (o funzione simile), appena dopo aver letto il payload:
    if payload.contacts:
        for c in payload.contacts:
            if getattr(c, "company_id", None):
                c.company_id = _normalize_name(c.company_id)
            if getattr(c, "company", None) and not getattr(c, "company_id", None):
                # fallback: usa company come id se manca company_id
                c.company_id = _normalize_name(c.company)


    # 1) corpus sorgente (riuso motore di ricerca centralizzato)
    ropt = payload.research or ResearchOptions()
    sources: List[dict] = _perform_research(payload) if ropt.mode != "off" else []

    # 2) split per competitor
    corpus_map: Dict[str, str] = {}
    if sources:
        by_comp = _split_corpus_by_competitor(payload.competitors or [], sources)
        for comp, items in by_comp.items():
            text = " ".join([(i.get("title","")+" "+i.get("snippet","")+" "+i.get("url","")).lower() for i in items])
            corpus_map[comp.lower()] = text

    # 3) normalizza lista feature
    normalized_features = _norm_features(payload)

    # 4) coverage feature×competitor (0..1) con “bonus” regex
    feature_coverage = {f: {} for f in normalized_features}
    for comp in (payload.competitors or []):
        text_corpus = corpus_map.get(comp.lower(), "")
        for feat in normalized_features:
            base = _score_feature_coverage(feat, corpus=text_corpus)    # tua logica base
            bonus = 0.25 * _feature_hit_score(feat, text_corpus)        # bonus regex (facoltativo)
            feature_coverage[feat][comp] = min(1.0, base + bonus)

    # 5) ranking (pesi industry/icp)
    ranking = _rank_competitors(
        feature_coverage=feature_coverage,
        industry=payload.industry,
        icp=payload.icp,
        corpus_map=corpus_map
    )

    # 6) research summary opzionale (LLM)
    research_summary = None
    if os.getenv("LLM_SUMMARY_ENABLED","false").lower() in {"1","true","yes"} and sources:
        try:
            research_summary = _summarize_corpus_items(sources, lang=(payload.language or "it"))
        except Exception:
            research_summary = None

    # 7) costruzione items e scoring per-item
    items: List[CompetitorItem] = []
    for comp in (payload.competitors or []):
        cov = {feat: feature_coverage[feat].get(comp, 0.0) for feat in normalized_features}
        item = CompetitorItem(name=comp, feature_coverage=cov)
        # arricchisci score_* per item in base al corpus
        item = _score_competitor_item(item, normalized_features, corpus_map.get(comp.lower(), ""), payload.industry, payload.icp)
        items.append(item)

    # 8) punteggi aggregati
    weights = (payload.meta or {}).get("feature_weights", {}) if getattr(payload, "meta", None) else {}
    for it in items:
        cov = (it.feature_coverage or {})
        it.score_feature = round(sum(cov.values()) / max(1, len(cov)), 4)
        it.score_total = _compute_item_total_score(it, features=list(normalized_features or []), weights=weights)

    # 9) ritorno matrice (ranking dentro .scoring)
    return CompetitorMatrix(
        company=payload.company,
        target=payload.company,
        industry=payload.industry,
        icp=payload.icp,
        normalized_features=normalized_features,
        feature_coverage=feature_coverage,
        items=items,
        research_summary=research_summary,
        scoring={
            "ranking": ranking,
            # pesi “macro” esposti per la UI/consumi a valle (come avevi prima)
            "weights": {"feature": 0.7, "industry_icp": 0.3},
            # se vuoi anche mostrare i pesi di dettaglio per feature, includi il dict usato su _compute_item_total_score:
            "feature_weights": weights,  # <- da (payload.meta or {}).get("feature_weights", {})
        },
    )

def rank_combinations(combos: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    # es: pesi su presenza trigger+allineamento persona
    for c in combos:
        c["score"] = round( 0.6*c.get("trigger_score",0) + 0.4*c.get("persona_fit",0), 3)
    return sorted(combos, key=lambda x: x["score"], reverse=True)


def _rank_competitors(
    feature_coverage: Dict[str, Dict[str, float]],
    industry: Optional[str] = None,
    icp: Optional[str] = None,
    weights: Optional[Dict[str, float]] = None,
    corpus_map: Optional[Dict[str, str]] = None

) -> List[Dict[str, Any]]:
    """
    Converte la matrice feature_coverage in un ranking:
    - somma pesata per competitor (peso=weights.get(feature,1.0))
    - boost semplice se il competitor compare spesso con industry/icp nel corpus concatenato (se disponibile)
    """
    weights = weights or {}
    # competitor list
    competitors = sorted({c for m in feature_coverage.values() for c in m.keys()})

    # totali per competitor
    rows = []
    for comp in competitors:
        tot = 0.0
        for feat, mp in feature_coverage.items():
            w = float(weights.get(feat, 1.0))
            v = float(mp.get(comp, 0.0))
            tot += w * v
        if corpus_map and (industry or icp):
            text = (corpus_map.get(comp.lower(), "") or "").lower()
            if industry and industry.lower() in text:
                tot += 0.05
            if icp and icp.lower() in text:
                tot += 0.05
        rows.append({"competitor": comp, "score": round(tot, 3)})

    rows.sort(key=lambda x: x["score"], reverse=True)
    return rows


def _compute_item_total_score(item, features: List[str], weights: Dict[str, float]) -> float:
    tot = 0.0
    for f in (features or []):
        w = float(weights.get(f, 1.0))
        v = float(item.feature_coverage.get(f, 0.0))
        tot += w * v
    return round(tot, 4)

def _matrix_to_feature_table_rows(matrix: "CompetitorMatrix") -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    feats = list(matrix.normalized_features or [])
    if not feats:
        bag = set()
        for it in (matrix.items or []):
            for f in (it.feature_coverage or {}).keys():
                bag.add(f)
        feats = sorted(bag)
    for f in feats:
        row = {"feature": f}
        for it in (matrix.items or []):
            row[it.name] = round(float(it.feature_coverage.get(f, 0.0)), 3)
        rows.append(row)
    return rows

def _rows_to_csv(rows: List[Dict[str, Any]]) -> str:
    if not rows: return ""
    buf = io.StringIO()
    w = csv.DictWriter(buf, fieldnames=list(rows[0].keys()))
    w.writeheader()
    for r in rows: w.writerow(r)
    return buf.getvalue()

def _business_day_add(d: date, n: int) -> date:
    # naive: aggiunge n giorni saltando sab/dom
    step = 1 if n >= 0 else -1
    k = abs(n)
    cur = d
    while k > 0:
        cur = cur + timedelta(days=step)
        if cur.weekday() < 5:
            k -= 1
    return cur

def _company_query(payload: GenerateSequenceRequest) -> str:
    c = (payload.contacts and payload.contacts[0].company) or ""
    return f"{c} EDI API news" if c else "EDI onboarding retail"

def next_workday(d: date) -> date:
    """Rimanda a lunedì se la data cade nel weekend, senza saltare giorni feriali."""
    while d.weekday() >= 5:  # 5=Sab, 6=Dom
        d += timedelta(days=1)
    return d

def parse_resource(file_path: str) -> str:
    file_path = file_path or ""
    fp = file_path.lower()
    if fp.endswith(".pdf"):
        return _extract_text_from_pdf(file_path)
    if fp.endswith((".png", ".jpg", ".jpeg")):
        # prova prima PIL+pytesseract; se fallisce prova opencv se disponibile
        try:
            from PIL import Image
            import pytesseract as _pt
            return _pt.image_to_string(Image.open(file_path)) or ""
        except Exception:
            try:
                import cv2  # type: ignore
                import pytesseract as _pt
                img = cv2.imread(file_path)
                return _pt.image_to_string(img) if img is not None else ""
            except Exception:
                return ""
    if fp.endswith((".docx",)):
        return _docx_to_text(file_path)
    # testo grezzo
    try:
        return Path(file_path).read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""

def _extract_text_from_pdf(path: str) -> str:
    try:
        from pdfminer.high_level import extract_text
        return extract_text(path) or ""
    except Exception:
        return ""

def _extract_text_from_image(path: str) -> str:

    """
    Best-effort: PIL + pytesseract (se presenti). Nessuna dipendenza hard.
    Niente cv2 per non avere 'cv2 not defined'.
    """
    try:
        from PIL import Image  # type: ignore
        import pytesseract  # type: ignore
        img = Image.open(path)
        return pytesseract.image_to_string(img) or ""
    except Exception:
        return

_URL_RE = re.compile(r'https?://[^\s)>\]"]+')

# ===================== BEGIN PATCH: KB search cache + TF-IDF mixed =====================
from pathlib import Path
import json, hashlib, time as _time

# Cache dirs
def _kbq_path(q: str) -> Path:
    h = hashlib.sha256((q or "").encode("utf-8")).hexdigest()
    return KB_QCACHE_DIR / f"{h}.json"

def _kbq_get(q: str):
    p = _kbq_path(q)
    if not p.exists(): return None
    try:
        obj = json.loads(p.read_text(encoding="utf-8"))
        if (_time.time() - float(obj.get("_ts", 0))) > KB_QCACHE_TTL:
            p.unlink(missing_ok=True); return None
        return obj.get("hits")
    except Exception:
        return None

def _kbq_put(q: str, hits: list[dict]):
    p = _kbq_path(q)
    try:
        p.write_text(json.dumps({"_ts": _time.time(), "hits": hits}, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass

# ---- TF-IDF helpers (soft-optional) ----
try:
    from sklearn.feature_extraction.text import TfidfVectorizer   # l’hai già in ambiente
    from sklearn.metrics.pairwise import cosine_similarity
except Exception:
    TfidfVectorizer = None
    cosine_similarity = None

def _do_kb_search_countbased(query: str, kb_items: list[dict], top_k: int = 5) -> list[dict]:
    # versione “semplice” che già usavi: tokenizzazione semplice + count + snippet + primo URL
    q = (query or "").lower().strip()
    out = []
    for it in kb_items:
        text = " ".join([it.get("title",""), it.get("snippet",""), it.get("url","")]).lower()
        cnt = text.count(q) if q else 0
        if cnt > 0:
            out.append({
                "title": it.get("title"),
                "url": it.get("url"),
                "snippet": it.get("snippet"),
                "score_count": float(cnt)
            })
    out.sort(key=lambda x: x["score_count"], reverse=True)
    return out[:top_k]

def _do_kb_search_tfidf(query: str, kb_items: list[dict], top_k: int = 5) -> list[dict]:
    if not (TfidfVectorizer and cosine_similarity):
        return []
    q = (query or "").strip()
    docs = [" ".join([it.get("title",""), it.get("snippet",""), it.get("url","")]) for it in kb_items]
    if not docs or not q: return []
    vec = TfidfVectorizer(ngram_range=(1,2), min_df=1, max_df=0.9)
    X = vec.fit_transform(docs + [q])
    sims = cosine_similarity(X[-1], X[:-1]).ravel()
    rows = []
    for i, it in enumerate(kb_items):
        if sims[i] <= 0: continue
        rows.append({
            "title": it.get("title"),
            "url": it.get("url"),
            "snippet": it.get("snippet"),
            "score_sim": float(sims[i])
        })
    rows.sort(key=lambda x: x["score_sim"], reverse=True)
    return rows[:top_k]

def _do_kb_search(query: str, top_k: int = 5) -> list[dict]:
    """
    Misto: 0.7 * count + 0.3 * similarity (TF-IDF)
    Mantiene snippet/primo URL come nella versione originale.
    Usa cache su disco per evitare ricerche duplicate.
    """
    if not query:
        return []

    # cache
    cached = _kbq_get(query)
    if cached is not None:
        return cached[:top_k]

    # 1) recupera lista KB “flat” (riusa la tua funzione/indice esistente)
    kb_items = get_kb_items_flat() # implementazione già presente nel file o equivalente

    # 2) calcola due ranking
    ca = _do_kb_search_countbased(query, kb_items, top_k=max(50, top_k))
    tf = _do_kb_search_tfidf(query, kb_items, top_k=max(50, top_k))

    # 3) merge per URL+title
    by_key = {}
    for r in ca:
        k = (r.get("title"), r.get("url"))
        by_key[k] = {"title": r.get("title"), "url": r.get("url"), "snippet": r.get("snippet"),
                     "score_count": float(r.get("score_count", 0.0)), "score_sim": 0.0}
    for r in tf:
        k = (r.get("title"), r.get("url"))
        if k not in by_key:
            by_key[k] = {"title": r.get("title"), "url": r.get("url"), "snippet": r.get("snippet"),
                         "score_count": 0.0, "score_sim": float(r.get("score_sim", 0.0))}
        else:
            by_key[k]["score_sim"] = float(r.get("score_sim", 0.0))

    # 4) ranking misto
    rows = []
    for _, r in by_key.items():
        score = 0.7 * min(1.0, r["score_count"]) + 0.3 * r["score_sim"]
        if score <= 0: continue
        r["score"] = round(float(score), 4)
        rows.append(r)

    rows.sort(key=lambda x: x["score"], reverse=True)
    rows = rows[:top_k]

    # cache write
    _kbq_put(query, rows)
    return rows
# ===================== END PATCH: KB search cache + TF-IDF mixed =====================


def _docx_from_messages(messages: List[Message], title: str = "Flowagent Sequence") -> Optional[bytes]:
    if Document is None:
        return None
    doc = Document()
    doc.add_heading(title, level=1)
    for m in messages or []:
        h = f"[{m.step}] {m.channel.upper()}  —  Variant: {getattr(m, 'variant', 'A')}"
        if getattr(m, "subject", None):
            h += f"  —  Subject: {m.subject}"
        doc.add_heading(h, level=2)
        p = doc.add_paragraph(m.text or "")
        p.style.font.size = Pt(11)
        if m.tips:
            doc.add_paragraph(f"Tips: {', '.join(m.tips)}").italic = True
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def export_messages_to_word(messages, path="messages.docx"):
    doc = Document()
    for m in messages:
        doc.add_heading(f"{m.step} - {m.channel}", level=2)
        doc.add_paragraph(f"Subject: {m.subject}")
        doc.add_paragraph(m.text)
        doc.add_paragraph(f"Tips: {', '.join(m.tips or [])}")
    doc.save(path)
    return path

def compose_thread_sequence(
    thread: ThreadSpec,
    payload: GenerateSequenceRequest,
    start_date: date,
    per_contact: bool = False,
    best_combo: dict | None = None,          # <-- passa "best" da generate_sequence
    research_meta: dict | None = None
) -> Tuple[List[Message], List[SequenceAction], List[CalendarEvent], dict]:  # ★ ritorna anche log
    messages: List[Message] = []
    actions: List[SequenceAction] = []
    calendar: List[CalendarEvent] = []
    # --- LOG strutturato ---
    log: dict = {
        "meta": {
            "start_date": str(start_date),
            "sequence_type": getattr(payload, "sequence_type", None),
            "language": getattr(payload, "language", "it"),
            "per_contact": bool(per_contact),
        },
        "persona": {},
        "research": {
            # se generate_sequence ti passa i meta, li appendi qui
            **(research_meta or {})
        },
        "ranking": {
            "best": best_combo  # <-- lo inseriamo qui, senza ricalcolare
        },
        "debug": {}
    }

    info = _get_or_synth_persona(
        getattr(payload, "industry", None),
        getattr(payload, "role", None),
        lang=getattr(payload, "language", "it")
    )

    # persona fallback deterministico + conf (TUO CODICE)
    persona_id = info["id"]

    # dopo il calcolo persona_id / fallback:
    persona_id = info["id"]
    log["persona"]["id"] = persona_id
    if info.get("fallback"):
        log["persona"]["fallback_note"] = "persona_fallback:" + info.get("why", "")
    if info.get("confidence") is not None:
        log["persona"]["fallback_confidence"] = info["confidence"]

    # quando ottieni research_hits (se li usi nella compose):
    # meta ricerca passati dall’alto (se presenti)
    if research_meta:
        try:
            log["research"]["hits_count"] = int(research_meta.get("hits_count", 0))
            if "mode" in research_meta:
                log["research"]["mode"] = research_meta["mode"]
        except Exception:
            pass

    # best combo già calcolata in generate_sequence
    log["ranking"]["best"] = best_combo


    channels: List[str] = thread.channels or ["email", "linkedin_dm", "inmail"]
    cadence = thread.cadence or _default_cadence_for(getattr(payload, "sequence_type", None))
    contacts = payload.contacts or []

    def _compose_subject_and_text(ch: str, c: Optional["Contact"]) -> Tuple[Optional[str], str]:
        if ch == "email":
            subj = "Onboarding EDI/API — idea veloce?"
            base = (
                f"Gentile {(getattr(c, 'first_name', '') or getattr(c, 'name', '')).strip() or ''},\n"
                "ho notato segnali di upgrade EDI; spesso l’impatto è su ticket e tempi di integrazione.\n"
                "Caso analogo: -35% ticket EDI e lead-time -50% in 90 giorni.\n"
                "Ha senso un confronto di 12’ per capire fit e tempi?\n"
            )
            return subj, base
        if ch in ("linkedin_dm", "inmail"):
            return None, "Spunto rapido su EDI/API – se utile 10’ questa settimana per capire fit e tempi."
        if ch == "call":
            return None, "Call breve per verificare fit e impatto sull’onboarding partner."
        return None, "Nota informativa sul tema EDI/API."

    def _step_day(i: int) -> int:
        try:
            if isinstance(cadence, list) and i < len(cadence) and isinstance(cadence[i], dict):
                return int(cadence[i].get("day", i * 3))
        except Exception:
            pass
        return i * 3

    def _mutate_B(subject: Optional[str], body: str, contact) -> tuple[Optional[str], str]:
        # (TUO CODICE)
        sB = subject
        if subject:
            s = subject.strip()
            if len(s) > 60:
                sB = s[:57] + "…"
            elif "—" not in s and "➜" not in s:
                sB = s + " — checklist 12’"
            else:
                sB = s
        else:
            sB = subject

        name = getattr(contact, "name", None) or getattr(contact, "first_name", "") or ""
        sal = f"Gentile {name.split()[0]}," if name else "Gentile,"
        lines = [l.strip() for l in (body or "").splitlines() if l.strip()]
        if lines:
            if len(lines) >= 2:
                lines[1] = ("In scenari analoghi abbiamo osservato una riduzione dei ticket e dei tempi "
                            "di integrazione già nel primo sprint.")
            lines.append("Se di interesse, posso inviare subito la checklist (12’). Le torna comodo mercoledì o giovedì?")
        bB = sal + "\n" + "\n".join(lines)
        return sB, bB

    # === ★ 1) TRIGGERS / PERSONA FIT / RESEARCH + BEST =======================
    # triggers semplici: usa payload.triggers se presente, altrimenti estrai da industry/role
    triggers: List[str] = []
    if getattr(payload, "industry", None):
        triggers.append(f"industry:{payload.industry.lower()}")
    if getattr(payload, "role", None):
        triggers.append(f"role:{payload.role.lower()}")
    if getattr(payload, "triggers", None):
        # se hai una struttura complessa, appiattisci in lista
        t = payload.triggers
        for k in ["erp", "competitor", "personal"]:
            for v in (getattr(t, k) or []):
                if v:
                    key = f"{k}:{v}".lower()
                    triggers.append(key)
    info = _get_or_synth_persona(getattr(payload, "industry", None),
                                getattr(payload, "role", None),
                                lang=getattr(payload, "language", "it"))
    persona_id = info["id"]
    fallback_note = ("persona_fallback:" + info["why"]) if info.get("fallback") else None
    fallback_confidence = info.get("confidence")

    persona_fit = persona_lookup_advanced(triggers, top_k=3)  # già nel file patchato. :contentReference[oaicite:1]{index=1}
    research_hits = _perform_research(payload)                # centralizzato su research_agent. :contentReference[oaicite:2]{index=2}
    best = pick_best_combo(triggers, persona_fit, research_hits)  # KPI/trigger scelti. :contentReference[oaicite:3]{index=3}

    # logging (se hai un dict log nello scope)
    if isinstance(log, dict):
        log["best_combo"] = {
            "persona_id": best.get("persona_id"),
            "trigger": best.get("trigger"),
            "score": best.get("score"),
            "evidence_count": len(research_hits or [])
        }
    # opzionale: etichette verso UI/StandardOutput
    labels = [f"persona:{best.get('persona_id','')}", f"trigger:{best.get('trigger','')}"]

    # === ★ 2) COSTRUZIONE MESSAGGI/AZIONI come prima, ma usando best =========
    step_idx = 0
    for ch in channels:
        step_idx += 1
        step = f"step-{step_idx:02d}"
        day_offset = _step_day(step_idx - 1)
        for c in (contacts or [None]):
            subj, body = _compose_subject_and_text(ch, c)

            # etichette con info "best"
            labels = [f"persona:{best.get('persona_id','')}", f"trigger:{best.get('trigger','')}"]
            # CTA/Openers safe con seed deterministico
            seed = _seed_from((subj or "") + (body or ""))
            mA = Message(channel=ch, step=step, variant="A", subject=subj, text=body, tips=None)
            mA = _apply_openers_cta_safe(mA, c, seed)  # ★ usa la safe ovunque

            # Variante B (mutata)
            sB, bB = _mutate_B(subj, body, c)
            mB = Message(channel=ch, step=step, variant="B", subject=sB, text=bB, tips=None)
            mB = _apply_openers_cta_safe(mB, c, seed + 7)

            # eventuale annotazione fallback persona
            if fallback_note:
                for m in (mA, mB):
                    m.tips = (m.tips or []) + [fallback_note, f"persona_fallback_conf:{fallback_confidence:.2f}"]

            messages.extend([mA, mB])

        # azione e calendario
        actions.append(SequenceAction(channel=ch, step=step, action="send", date=start_date + timedelta(days=day_offset)))
        calendar.append(CalendarEvent(date=str(start_date + timedelta(days=day_offset)), action="send", channel=ch, step=step))

    # pianificatori LI (bump/followup) – riusa i tuoi (già nel file)
    plan_inmail_bump_if_no_connection(messages, actions, calendar, thread, payload, start_date)  # :contentReference[oaicite:4]{index=4}
    plan_dm_followup_on_connection_accept(messages, actions, calendar, thread, payload, start_date)  # :contentReference[oaicite:5]{index=5}

    # === ★ 3) LOG COMPLETO + CANONICAL ======================================
    # accumulatori prima del loop
    all_msgs: list[Message] = []
    all_actions: list[SequenceAction] = []
    all_calendar: list[CalendarEvent] = []
    logs: list[dict] = []

    msgs, acts, cal, log = compose_thread_sequence(
    thread, payload, start_date=date.today(),
    per_contact=False,
    best_combo=best,
    research_meta={
        "hits_count": len(research_hits or []),
        "mode": getattr(getattr(payload, "research", None), "mode", "kb_only"),
    }
)


    if msgs: all_msgs += msgs
    if acts: all_actions += acts
    if cal:  all_calendar += cal
    if log:  logs.append(log)

    if msgs: all_msgs += msgs
    if acts: all_actions += acts
    if cal:  all_calendar += cal
    if log:  logs.append(log)

    log = _drop_nones(log or {})
    log.update({
        "persona_id": persona_id,
        "triggers": triggers,
        "persona_fit": persona_fit[:3] if persona_fit else [],
        "best_combo": best,
        "research_hits": len(research_hits or []),
    })

    try:
        log["_canonical"] = _json_canonical(log)
    except Exception:
        pass
    output_labels = (output_labels or []) + labels
    labels = [f"persona:{best.get('persona_id','')}", f"trigger:{best.get('trigger','')}"]
    return messages, actions, calendar, log

def _load_campaigns() -> List[Dict[str, Any]]:
    """
    Carica campagne da:
      - ./campaigns/*.csv (header libero, unisci tutte le colonne)
      - ./campaigns/campaigns.json (lista di oggetti)
    Ordina per data se presente un campo 'date' ISO-like.
    """
    data_dir = Path("./data"); data_dir.mkdir(parents=True, exist_ok=True)
    csv_path = data_dir / "campaigns.csv"
    json_path = data_dir / "campaigns.json"

    rows: List[Dict[str, Any]] = []
    if csv_path.exists():
        import csv
        with csv_path.open("r", encoding="utf-8", newline="") as f:
            reader = csv.DictReader(f)
            for r in reader:
                rows.append(dict(r))
    elif json_path.exists():
        obj = json.loads(json_path.read_text(encoding="utf-8"))
        if isinstance(obj, list):
            rows = [dict(x) for x in obj if isinstance(x, dict)]

    # normalizzazioni minime
    for r in rows:
        r.setdefault("date", datetime.utcnow().strftime("%Y-%m-%d"))
        r.setdefault("contact", "")
        r.setdefault("action", "")

    def _key(r):
        d = r.get("date") or r.get("Date") or r.get("created_at")
        try:
            return datetime.fromisoformat(str(d).replace("Z","").strip())
        except Exception:
            return datetime.min
    rows.sort(key=_key, reverse=True)
    return rows


def _dm_mapping(industry: Optional[str], product: Optional[str]) -> List[DmRole]:
    base = [
        ("CIO / IT Director",
         ["Affidabilità integrazioni", "Sicurezza dei dati", "Effort IT"],
         ["Ticket rate", "Lead-time integrazioni", "MTTR"],
         "Riduci effort IT e ticket con connettori pre-validati"),
        ("Operations / Supply Chain",
         ["Onboarding partner", "SLAs con fornitori", "Visibilità flussi"],
         ["Tempo onboarding", "% partner attivi", "Errori documentali"],
         "Onboarding partner più rapido e tracciabile"),
        ("Finance / CFO",
         ["Costo totale", "Rischio errori", "Cash conversion"],
         ["Costo integrazione/partner", "Chargebacks", "DSO"],
         "ROI in 90 giorni su costi e chargebacks"),
    ]
    return [DmRole(role=r, cares_about=ca, kpis=k, angle=ang) for r, ca, k, ang in base]


def _price_objection_responses(product: str, context: Optional[str], currency: str) -> List[str]:
    ctx = f" ({context})" if context else ""
    r1 = (f"Capisco il tema budget. Il rischio è il costo dell’inazione{ctx}: ticket e ritardi EDI continuano a "
          f"impattare SLA e chargebacks. Con {product} tipicamente riduciamo i ticket ~30–40% in 90 giorni: "
          f"l’impatto economico supera la differenza di prezzo.")
    r2 = (f"Se guardiamo 12 mesi, {product} incide su onboarding partner e lead time. Anche un miglioramento "
          f"conservativo del 20–30% libera ore IT e riduce penali: il TCO scende e il payback è < 1Q.")
    r3 = (f"Possiamo iniziare con uno scope più snello per restare nel budget, misurando ROI su 1–2 aree critiche; "
          f"poi estendiamo. Così riduciamo il rischio e massimizziamo il valore.")
    return [r1, r2, r3]


PERSONA_INDEX = {}

# --- ADD HERE: Persona synth fallback (app.py) ---
PERSONAS_GEN_DIR = Path("kb/personas/generated")
PERSONAS_GEN_DIR.mkdir(parents=True, exist_ok=True)

def _slugify(s: str) -> str:
    s = (s or "").lower()
    s = re.sub(r"[^a-z0-9]+", "-", s).strip("-")
    return s or "persona"

def _synth_persona(industry: str, role: str, lang: str = "it") -> dict:
    # versione deterministica + hook LLM opzionale
    base = {
        "id": _slugify(f"{industry}-{role}"),
        "title": f"{role} in {industry}",
        "lang": lang,
        "kpis": ["lead_time", "onboarding", "compliance", "integration"],
        "pains": ["ritardi integrazione", "ticket elevati", "mancata visibilità", "costi di change"],
        "tone": "conciso, autorevole, numerico",
        "framework_pref": ["TIPPS","POKE"],
    }
    if os.getenv("LLM_SYNTH_PERSONA_ENABLED","false").lower() in {"1","true","yes"} and os.getenv("OPENAI_API_KEY"):
        try:
            import openai  # se presente
            openai.api_key = os.getenv("OPENAI_API_KEY")
            prompt = f"Crea una buyer persona sintetica (JSON) per ruolo '{role}' nel settore '{industry}', lingua '{lang}'. Campi: kpis, pains, triggers, objections, tone, messages_notes."
            model = os.getenv("OPENAI_MODEL","gpt-4o-mini")
            rsp = openai.ChatCompletion.create(
                model=model,
                messages=[{"role":"system","content":"Sei un analista B2B."},
                          {"role":"user","content":prompt}]
            )
            js = rsp.choices[0].message["content"]
            enriched = json.loads(js)
            base.update({k: v for k, v in enriched.items() if k in {"kpis","pains","triggers","objections","tone","messaging_notes"}})
        except Exception:
            pass
    (PERSONAS_GEN_DIR / f"{base['id']}.json").write_text(json.dumps(base, ensure_ascii=False, indent=2), encoding="utf-8")
    return base

def _get_or_synth_persona(industry: str, role: str, lang: str = "it") -> dict:
    pid, why, conf = _closest_persona_id(industry, role)
    if pid and pid.lower() not in {"generic","unknown"}:
        return {"id": pid, "fallback": False, "confidence": conf, "why": why}
    # synth nuova persona
    persona = _synth_persona(industry or "generic", role or "generic", lang=lang)
    return {"id": persona["id"], "fallback": True, "confidence": 0.51, "why": "synth"}


def _list_known_personas() -> List[Tuple[str, str, str]]:
    """
    Ritorna [(persona_id, industry, role), ...] dalla tua libreria.
    Sostituisci/aggancia alla tua sorgente; default minimale.
    """
    out = []
    for pid, meta in (PERSONA_INDEX or {}).items():  # se hai già un indice
        out.append((pid, meta.get("industry",""), meta.get("role","")))
    return out

def _normalize_research_corpus(hits: List[dict]) -> List[dict]:
    """
    Normalizza i risultati di ricerca in un formato coerente:
    - title, snippet, url, domain, text (title+snippet)
    """
    norm = []
    for h in hits or []:
        title = (h.get("title") or "").strip()
        snip  = (h.get("snippet") or "").strip()
        url   = (h.get("url") or "").strip()
        dom   = ""
        try:
            dom = urlparse(url).netloc.lower()
        except Exception:
            dom = ""
        text = (title + " " + snip).strip()
        norm.append({
            "title": title,
            "snippet": snip,
            "url": url,
            "domain": dom,
            "text": text
        })
    return norm


def _summarize_corpus_items(items: List[Dict[str, Any]], lang: str = "it") -> str:
    """
    Wrapper: prende items (title/url/snippet) e delega a _summarize_sources_llm.

    """
    return _summarize_sources_llm(items or [], lang=lang)

def _default_cadence_for(sequence_type: str | None) -> list[dict]:
    """
    Ritorna una cadence di default in assenza di thread.cadence.
    sequence_type: "with_inmail", "without_inmail" o None
    """
    if (sequence_type or "").lower() == "with_inmail":
        return [
            {"day": 0, "action": "email"},
            {"day": 3, "action": "linkedin_dm"},
            {"day": 7, "action": "inmail"},
        ]
    else:
        return [
            {"day": 0, "action": "email"},
            {"day": 3, "action": "linkedin_dm"},
            {"day": 7, "action": "email_followup"},
        ]

def _llm_complete(system: str, user: str, model: Optional[str] = None, temperature: float = 0.2) -> str:
    """
    Wrapper LLM con OpenAI (se disponibile) o fallback deterministico.
    Richiede env OPENAI_API_KEY e (opzionale) MODEL_LLM (es: gpt-4o-mini).
    """
    try:
        from openai import OpenAI  # type: ignore
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError("OPENAI_API_KEY assente")
        client = OpenAI(api_key=api_key)
        mdl = model or os.getenv("MODEL_LLM", "gpt-4o-mini")
        resp = client.chat.completions.create(
            model=mdl,
            temperature=temperature,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        # Fallback deterministico minimale
        return f"""Persona (fallback deterministico):
- Focus: {user}
- Pains: integrazione lenta, ticket elevati, time-to-value lungo
- Symptoms: ritardi onboarding partner, colli di bottiglia IT
- KPIs: lead time onboarding, tasso errori EDI, ticket/mese
- Objections: complessità progetto, tempo IT, budget
- Hooks: template pre-validati, ROI 90 giorni, riduzione ticket
"""

def _normalize_label(s: Optional[str]) -> str:
    return (s or "").strip().lower()

def _merge_lists_unique(*lists: List[str]) -> List[str]:
    seen, out = set(), []
    for L in lists:
        for x in L or []:
            k = _normalize_label(x)
            if k and k not in seen:
                seen.add(k); out.append(x)
    return out

def _buyerp_from_doc(doc: Any) -> Dict[str, Any]:
    """Normalizza un record persona già caricato in indice in un dizionario coerente."""
    return {
        "persona_id": getattr(doc, "persona_id", None) or getattr(doc, "id", None),
        "industry": getattr(doc, "industry", None),
        "role": getattr(doc, "role", None),
        "pains": getattr(doc, "pains", None) or getattr(doc, "pain", None) or [],
        "symptoms": getattr(doc, "symptoms", None) or getattr(doc, "symptom", None) or [],
        "kpis": getattr(doc, "kpis", None) or getattr(doc, "kpi", None) or [],
        "objections": getattr(doc, "objections", None) or [],
        "hooks": getattr(doc, "hooks", None) or getattr(doc, "messages", None) or [],
        "priorities": getattr(doc, "priorities", None) or [],
    }

# ====== BEGIN: creative persona synth ======

LLM_SYNTH_PERSONA_ENABLED = os.getenv("LLM_SYNTH_PERSONA_ENABLED","false").lower() in {"1","true","yes"}

def synthesize_creative_persona(base_persona: dict, sector: str, language: str="it") -> dict:
    """
    base_persona: tua persona 'classica' (pain, KPI, objections, goals...)
    Ritorna: persona 'creative' con campi extra per stile/voce/hook
    """
    if not LLM_SYNTH_PERSONA_ENABLED:
        return {**base_persona, "creative": {"enabled": False, "note": "disabled"}}

    # usa il tuo helper interno (es. _llm_complete / _llm_chat / _polish)
    prompt = f"""
    Sei un copy strategist B2B. Trasforma questa buyer persona in una 'Creative Persona' per contenuti TIPPS/LinkedIn.
    Output in JSON MINIFICATO.
    CONTENUTO BASE (settore: {sector}, lingua: {language}):
    {base_persona}

    Struttura precisa:
    {{
      "tone": ["consultivo","pragmatico","anti-jargon"],
      "style_guidelines": {{
        "sentence_length": "varia 6-16 parole",
        "verbs": "verbi d'azione, voce attiva",
        "jargon_to_avoid": ["trasformazione olistica","sinergie end-to-end"]
      }},
      "hooks": {{
        "status_quo": ["E se il problema non fosse l’ERP ma il time-to-integration?"],
        "contrast": ["2 ore per un partner vs 2 settimane"],
        "evidence": ["ASN accuracy +18% in 90 giorni"]
      }},
      "analogy_bank": ["control tower","tubi/valvole dei dati","supply chain come semafori"],
      "emotional_triggers": ["rischio reputazionale","SLA clienti","notte tranquilla"],
      "poke_questions": ["Cosa rende accettabile un 4% di IDoc falliti?"],
      "coi_prompts": ["Cosa costa NON ridurre del 30% l'onboarding dei partner?"],
      "lexicon": {{"preferred": ["integrazione","SLA","lead time"], "avoid": ["digitalizzazione olistica"]}}
    }}
    """
    js = _llm_complete(prompt, temperature=0.3, json_expected=True)  # riusa il tuo helper
    return {**base_persona, "creative": {"enabled": True, **js}}
# ====== END: creative persona synth ======


def _synth_fallback_from_corpus(industry: Optional[str], role: Optional[str]) -> Dict[str, Any]:
    """
    Costruisce una persona sintetica senza LLM, fondendo ciò che trova nell'indice
    (match per industry/role) + un baseline generico.
    """
    docs = _load_persona_index()  # deve già esistere nel tuo file
    ind = _normalize_label(industry)
    rol = _normalize_label(role)

    pool = []
    for d in docs:
        di = _normalize_label(d.industry)
        dr = _normalize_label(d.role)
        if (ind and di == ind) or (rol and dr == rol):
            pool.append(_buyerp_from_doc(d))

    # baseline molto generico
    base = {
        "persona_id": f"synth-{ind or 'any'}-{rol or 'any'}",
        "industry": industry,
        "role": role,
        "pains": ["integrazione lenta", "ticket elevati", "time-to-value lungo"],
        "symptoms": ["ritardi onboarding partner", "colli di bottiglia IT", "scarsa visibilità flussi"],
        "kpis": ["lead time onboarding", "error rate EDI", "ticket/mese", "tempo medio risoluzione"],
        "objections": ["complessità progetto", "tempo IT richiesto", "budget ridotto"],
        "hooks": ["template pre-validati", "ROI in 90 giorni", "riduzione ticket"],
        "priorities": ["stabilizzare EDI", "accelerare onboarding", "ridurre costi supporto"],
        "synthesized": True,
    }

    if not pool:
        return base

    # fondi pool in modo deterministico
    pains = _merge_lists_unique(*[p.get("pains", []) for p in pool], base["pains"])
    symptoms = _merge_lists_unique(*[p.get("symptoms", []) for p in pool], base["symptoms"])
    kpis = _merge_lists_unique(*[p.get("kpis", []) for p in pool], base["kpis"])
    objections = _merge_lists_unique(*[p.get("objections", []) for p in pool], base["objections"])
    hooks = _merge_lists_unique(*[p.get("hooks", []) for p in pool], base["hooks"])
    priorities = _merge_lists_unique(*[p.get("priorities", []) for p in pool], base["priorities"])

    base.update({
        "pains": pains[:12],
        "symptoms": symptoms[:12],
        "kpis": kpis[:12],
        "objections": objections[:12],
        "hooks": hooks[:12],
        "priorities": priorities[:12],
    })
    return base

def _synthesize_persona_with_llm(industry: Optional[str], role: Optional[str], lang: str = "it") -> Dict[str, Any]:
    """
    Restituisce un dict persona "synth". Se LLM abilitato -> usa LLM; altrimenti fallback deterministico da corpus.
    - Garantisce schema completo e normalizzato.
    - Taglia/deduplica le liste.
    - Genera persona_id "synth-{industry}-{role}" sanificato.
    """
    def _sanitize_id(x: str) -> str:
        return re.sub(r"[^a-z0-9\-]+", "-", (x or "").strip().lower())

    enabled = os.getenv("LLM_SYNTH_PERSONA_ENABLED", "false").lower() in {"1", "true", "yes"}
    if not enabled:
        return _synth_fallback_from_corpus(industry, role)

    system = "Sei un esperto di sales enablement. Genera profili buyer persona strutturati e coerenti."
    user = (
        f"Crea una buyer persona per:\n"
        f"- Industry: {industry or 'any'}\n"
        f"- Role: {role or 'any'}\n\n"
        "Rispondi in JSON con queste chiavi:\n"
        "persona_id, industry, role, pains[], symptoms[], kpis[], objections[], hooks[], priorities[].\n"
        "Voci concise, non generiche, max 12 elementi per lista."
    )
    try:
        raw = _llm_complete(system, user, lang=lang)  # se la tua _llm_complete accetta lang
    except Exception:
        raw = None

    if not raw:
        return _synth_fallback_from_corpus(industry, role)

    try:
        data = json.loads(raw)
    except Exception:
        return _synth_fallback_from_corpus(industry, role)

    # schema hardening
    persona_id = data.get("persona_id") or f"synth-{_sanitize_id(industry or 'any')}-{_sanitize_id(role or 'any')}"
    out = {
        "persona_id": persona_id,
        "industry": data.get("industry") or industry,
        "role": data.get("role") or role,
        "pains": list(dict.fromkeys([str(x).strip() for x in (data.get("pains") or [])]))[:12],
        "symptoms": list(dict.fromkeys([str(x).strip() for x in (data.get("symptoms") or [])]))[:12],
        "kpis": list(dict.fromkeys([str(x).strip() for x in (data.get("kpis") or [])]))[:12],
        "objections": list(dict.fromkeys([str(x).strip() for x in (data.get("objections") or [])]))[:12],
        "hooks": list(dict.fromkeys([str(x).strip() for x in (data.get("hooks") or [])]))[:12],
        "priorities": list(dict.fromkeys([str(x).strip() for x in (data.get("priorities") or [])]))[:12],
        "synthesized": True,
    }
    return out

def _string_sim(a: Optional[str], b: Optional[str]) -> float:
    a = (a or "").strip().lower()
    b = (b or "").strip().lower()
    if not a and not b:
        return 0.0
    return SequenceMatcher(None, a, b).ratio()

def _top_matches(corpus: str, candidates: List[str], top_k: int = 3) -> List[Dict[str, Any]]:
    """
    Ritorna i migliori match nel corpus contro una lista di candidati.
    Score = max(count-based, similarity).
    """
    text = (corpus or "").lower()
    out = []
    for c in candidates:
        c0 = c.lower()
        cnt = text.count(c0)
        sim = SequenceMatcher(None, text[: min(len(text), 2000)], c0).ratio()
        score = max(min(1.0, cnt / 3.0), sim)
        if score > 0:
            out.append({"value": c, "score": round(float(score), 3)})
    out.sort(key=lambda x: x["score"], reverse=True)
    return out[:top_k]

# TIPPS minimi: 2 bucket di tips usabili ovunque
tipsA = [
    "Anchora al contesto: cita 1 fatto oggettivo dall’azienda (fonte pubblica).",
    "Quantifica l’impatto (es. '−20% lead time integrazioni').",
    "Proponi un micro-impegno chiaro (12’ call).",
]
tipsB = [
    "Riduci il gergo tecnico: 1 concetto per frase.",
    "Evidenzia il 'perché ora' (trigger temporale).",
    "Chiudi con 1 domanda binaria (mercoledì o giovedì?).",
]

def _ics_escape(s: str) -> str:
    return (s or "").replace("\\", "\\\\").replace(";", "\\;").replace(",", "\\,").replace("\n", "\\n")

def _fmt_dt(dt: datetime) -> str:
    # sempre UTC nell’ICS
    return dt.astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")

def _ics_from_calendar(events: List["CalendarEvent"], cal_name: str = "FlowAgent v3") -> str:
    lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        f"X-WR-CALNAME:{_ics_escape(cal_name)}",
        "PRODID:-//FlowAgent//EN",
    ]
    for ev in events or []:
        # --- BLOCCO CHE NON SAPEVI DOVE METTERE (eccolo, qui dentro) ---
        start_dt = getattr(ev, "start_datetime", None)
        end_dt   = getattr(ev, "end_datetime", None)
        if not start_dt and getattr(ev, "date", None):
            start_dt = datetime.combine(ev.date, datetime.min.time()).astimezone(timezone.utc)
        if not end_dt:
            end_dt = (start_dt + timedelta(minutes=5)) if start_dt else datetime.utcnow() + timedelta(minutes=5)

        dtstart = _fmt_dt(start_dt)
        dtend   = _fmt_dt(end_dt)
        summary = _ics_escape(getattr(ev, "title", None) or getattr(ev, "action", None) or "Flowagent task")
        desc    = _ics_escape(getattr(ev, "description", None) or "")
        loc     = _ics_escape(getattr(ev, "location", None) or "")

        uid = f"{uuid.uuid4()}@flowagent"
        lines += [
            "BEGIN:VEVENT",
            f"UID:{uid}",
            f"DTSTART:{dtstart}",
            f"DTEND:{dtend}",
            f"SUMMARY:{summary}",
            f"DESCRIPTION:{desc}",
            f"LOCATION:{loc}",
            "END:VEVENT"
        ]
    lines.append("END:VCALENDAR")
    return "\r\n".join(lines)
# === OCR/PDF/IMG — estrattore unico "prod ready" ==============================
MAX_BYTES = int(os.getenv("INGEST_MAX_BYTES", str(15 * 1024 * 1024)))  # 15MB
ALLOWED_EXT = {".pdf", ".docx", ".txt", ".md", ".png", ".jpg", ".jpeg", ".tif", ".tiff"}

def _safe_read_bytes(path: Path, max_bytes: int = MAX_BYTES) -> bytes:
    if not path.exists() or not path.is_file():
        return b""
    if path.stat().st_size > max_bytes:
        # tronca in sicurezza in caso di PDF/scansioni enormi
        with path.open("rb") as f:
            return f.read(max_bytes)
    return path.read_bytes()

def extract_text_any(path: Path) -> str:
    """
    Estrae testo da: PDF, DOCX, TXT/MD, immagini (OCR).
    - PDF → pdfminer
    - DOCX → docx2txt
    - TXT/MD → lettura diretta
    - Immagini → PIL+Tesseract (se presenti)
    """
    try:
        ext = path.suffix.lower()
        if ext not in ALLOWED_EXT:
            return ""

        # 1) TXT / MD
        if ext in {".txt", ".md"}:
            try:
                return path.read_text(encoding="utf-8", errors="ignore")
            except Exception:
                raw = _safe_read_bytes(path)
                try:
                    return raw.decode("utf-8", errors="ignore")
                except Exception:
                    return ""

        # 2) DOCX
        if ext == ".docx" and docx2txt:
            try:
                return docx2txt.process(str(path)) or ""
            except Exception:
                pass  # fallback all’OCR sotto

        # 3) PDF
        if ext == ".pdf" and pdf_extract_text:
            try:
                # pdfminer è sicuro lato RCE; limitiamo comunque dimensione
                tmp = Path(str(path))
                return pdf_extract_text(str(tmp)) or ""
            except Exception:
                pass  # fallback all’OCR sotto

        # 4) Immagini / fallback OCR
        if (ext in {".png",".jpg",".jpeg",".tif",".tiff"} or True) and pytesseract:
            try:
                img_bytes = _safe_read_bytes(path)
                if not img_bytes:
                    return ""
                from PIL import Image
                import io as _io
                im = Image.open(_io.BytesIO(img_bytes))
                # normalizzazione semplice per OCR
                if im.mode not in ("L","RGB"):
                    im = im.convert("RGB")
                # denoise base se cv2 è disponibile
                try:
                    import cv2
                    import numpy as np
                    arr = np.array(im)
                    arr = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
                    arr = cv2.medianBlur(arr, 3)
                    im = Image.fromarray(arr)
                except Exception:
                    pass
                txt = pytesseract.image_to_string(im, lang=os.getenv("OCR_LANG","ita+eng"))
                return txt or ""
            except Exception:
                return ""
    except Exception:
        return ""
    return ""

KB_RAW = Path("kb/raw")
KB_RAW.mkdir(parents=True, exist_ok=True)
for p in KB_RAW.glob("**/*"):
    if not p.is_file():
        continue
    side = p.with_suffix(".txt")
    if side.exists():
        continue
    txt = extract_text_any(p)  # <— usa il nuovo estrattore robusto
    if txt.strip():
        side.write_text(txt, encoding="utf-8")


# === 7) VALORI DI DEFAULT “GLOBALI” USATI IN PIÙ PUNTI =========================

@app.get("/", include_in_schema=False)
def root_banner():
    app_name = "Flowagent V3 Orchestrator"
    app_ver  = os.getenv("APP_VERSION", "1.1.0")
    build_sha = os.getenv("BUILD_SHA", "render")
    app_env  = APP_ENV
    docs_url = "/docs" if docs_on else "/openapi.json"

    html = f"""
    <!doctype html>
    <html><head><meta charset="utf-8"/><meta name="viewport" content="width=device-width,initial-scale=1"/>
    <title>{app_name}</title>
    <style>body{{font-family:system-ui,-apple-system,Segoe UI,Roboto,Arial,sans-serif;margin:0;background:#0b1020;color:#e6e8ef}}
    .wrap{{max-width:820px;margin:6rem auto;padding:2rem;text-align:center}}
    .card{{background:#121936;border:1px solid #222a51;border-radius:14px;padding:2rem}}
    h1{{margin:0 0 1rem;font-size:1.8rem}}
    .meta{{opacity:.8;margin:.5rem 0 1.2rem}}
    .btns a{{display:inline-block;margin:.25rem;padding:.7rem 1rem;border-radius:10px;text-decoration:none;border:1px solid #3b4799}}
    .btns a.primary{{background:#2b3bbb;color:white;border-color:#2b3bbb}}
    .btns a:hover{{filter:brightness(1.08)}}</style></head>
    <body><div class="wrap"><div class="card">
    <h1>{app_name}</h1>
    <div class="meta">Env: <code>{app_env}</code> · Versione: <code>{app_ver}</code> · Build: <code>{build_sha[:7]}</code></div>
    <div class="btns"><a class="primary" href="{docs_url}">📚 API Docs</a>
    <a href="/openapi.json">🧩 OpenAPI</a> <a href="/health">✅ Health</a></div>
    </div></div></body></html>
    """
    return HTMLResponse(html)

@app.get("/ping", include_in_schema=False)
def ping():
    return {"service":"flowagent-v3","status":"ok","docs":"/docs"}

@app.get("/health")
def health():
    return {"ok": True, "time": datetime.utcnow().isoformat()+ "Z"}

Mode = Literal["AE", "SDR"]
Level = Literal["Beginner", "Intermediate", "Advanced"]
SequenceType = Literal["with_inmail", "without_inmail"]
Framework = Literal["Auto", "TIPPS", "TIPPS+COI", "Poke", "HarrisNEAT", "NEAT_structure", "ShowMeYouKnowMe", "BoldInsight", "Challenger"]
Channel = Literal["email", "linkedin_dm", "inmail", "voice_note", "video_dm", "call"]

@app.post("/diff/messages", response_model=DiffResponse)
def diff_messages(req: DiffRequest, authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    chunks=[]
    # accoppia per step+channel
    key = lambda m: (m.step or "", m.channel or "")
    amap = {key(m):m for m in (req.a or [])}
    bmap = {key(m):m for m in (req.b or [])}
    keys = sorted(set(list(amap.keys())+list(bmap.keys())))
    for k in keys:
        a = (amap.get(k).text or "").splitlines() if amap.get(k) else []
        b = (bmap.get(k).text or "").splitlines() if bmap.get(k) else []
        diff = "\n".join(difflib.unified_diff(a, b, fromfile="A", tofile="B", lineterm=""))
        chunks.append(DiffChunk(step=k[0], channel=k[1], diff_unified=diff))
    return DiffResponse(chunks=chunks)

fb_conf = None  # or fb_conf = 0.0

@app.post("/sequence/compose", response_model=StandardOutput)
def sequence_compose(req: GenerateSequenceRequest, authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    thread = req.threads[0] if (getattr(req, "threads", None)) else (req.thread or ThreadSpec(name="default"))

    result = compose_thread_sequence(thread, req, start_date=date.today(), per_contact=False)
    # compat 3-ple/4-ple
    if len(result) == 4:
        msgs, acts, cal, log = result
    else:
        msgs, acts, cal = result
        log = {}

    # pulizia log qui (ridondante ma ok)
    log = _drop_nones(log)
    try: log["_canonical"] = _json_canonical(log)
    except Exception: pass

    return StandardOutput(messages=msgs, actions=acts, calendar=cal, logging=log)


@app.post("/run/generate_sequence", response_model=StandardOutput)
def generate_sequence(
    payload: GenerateSequenceRequest,
    authorization: Optional[str] = Header(None),
    idem_hdr: Optional[str] = Header(default=None, alias="Idempotency-Key"),
    idem_q: Optional[str] = Query(default=None, alias="idempotency_key"),
    approved: Optional[str] = Header(default=None, alias=APPROVAL_HEADER),
):
    _guard_ip(json.dumps(payload.dict() if hasattr(payload, "dict") else payload))
    ensure_auth(authorization)

    idempotency_key = idem_hdr or idem_q
    if idempotency_key:
        prior = idem_get(idempotency_key)
        if prior is not None:
            return prior

    consequential = False
    approval_gate(consequential, approved)

    ropt = _effective_research_mode(payload.research or ResearchOptions())
    research_hits: List[dict] = []
    if ropt.mode != "off":
        try:
            research_hits = _perform_research(payload)  # usa cache
        except Exception:
            research_hits = []

    research_summary: Optional[str] = None
    if os.getenv("LLM_SUMMARY_ENABLED","false").lower() in {"1","true","yes"} and research_hits:
        try:
            research_summary = _summarize_sources_llm(
                research_hits, lang=(getattr(payload, "language", "it") or "it")
            )
        except Exception:
            research_summary = None

    try:
        rr_norm = _normalize_research_corpus(research_hits)
    except Exception:
        rr_norm = None

    agg_text = " ".join([(s.get("title","") + " " + s.get("snippet","")) for s in (research_hits or [])])
    triggers = _extract_triggers_from_text(agg_text)
    persona_fit = persona_lookup_advanced(triggers, top_k=3)
    best = pick_best_combo(triggers, persona_fit, research_hits)

    # 2) log/labels — INIZIALIZZA per evitare "not defined"
    logging: dict = {}
    labels: List[str] = []

    logging["best_combo"] = best
    if best and isinstance(best, dict):
        if best.get("trigger"):
            labels.append(f"TR:{best['trigger']}")
        if best.get("persona_id"):
            labels.append(f"PE:{best['persona_id']}")
    thread = payload.threads[0] if (getattr(payload, "threads", None)) else ThreadSpec(name="default")

    # 4) compongo una o più sequenze
    all_msgs: List[Message] = []
    all_actions: List[SequenceAction] = []
    all_calendar: List[CalendarEvent] = []
    all_logs: List[dict] = []

    # passaggio opzionale alla compose se vuoi usarlo nelle micro-scelte
    msgs, acts, cal, log = compose_thread_sequence(
        thread, payload, start_date=date.today(),
        per_contact=False,
        best_combo=best,
        research_meta={
            "hits_count": len(research_hits or []),
            "mode": getattr(getattr(payload, "research", None), "mode", "kb_only"),
        }
    )

    # 👉 Prima ACCUMULO…
    if msgs: all_msgs += msgs
    if acts: all_actions += acts
    if cal:  all_calendar += cal
    if log:  all_logs.append(log)

    # …POI leggo/derivo (il “Punto 6”)
    persona_conf = {}  # se ti serve per export/log
    persona_conf["total_msgs"] = len(all_msgs)
    persona_conf["actions"]    = len(all_actions)
    persona_conf["calendar"]   = len(all_calendar)

    # arricchisco log senza sovrascrivere
    log = _drop_nones(log or {})
    log.update({
        "triggers": triggers,
        "persona_fit": persona_fit[:3] if persona_fit else [],
        "best_combo": best,
        "research_hits": len(research_hits or []),
        "research_summary": research_summary,
        "metrics": persona_conf,
    })

    # === Persona Synthesis LLM (opzionale, dietro flag) ===
    if os.getenv("LLM_SYNTH_PERSONA_ENABLED", "false").lower() in {"1", "true", "yes"}:
        try:
            industry = getattr(payload, "industry", None)
            role = getattr(payload, "role", None)
            if industry or role:
                synth_text = _synthesize_persona_with_llm(industry, role)
                synth_id = f"synth-{(industry or 'any').lower()}-{(role or 'any').lower()}"
                synth_path = Path("kb/personas") / f"{synth_id}.json"
                synth_path.write_text(json.dumps({
                    "persona_id": synth_id,
                    "industry": industry,
                    "role": role,
                    "synthesized": True,
                    "content": synth_text
                }, ensure_ascii=False, indent=2))
                # aggiorna indice globale
                _refresh_persona_index()
                persona_synth_id = synth_id
        except Exception as e:
            persona_synth_error = str(e)
    # ---- compongo threads ----
    all_msgs: List[Message] = []; all_actions: List[SequenceAction] = []; all_calendar: List[CalendarEvent] = []
    persona_conf_samples: List[float] = []

    if payload.threads:
        for th in payload.threads:
            per_contact_flag = bool(getattr(th, "meta", {}).get("per_contact", False))
            if ALLOW_PER_CONTACT: per_contact_flag = True
            msgs, acts, cal = compose_thread_sequence(th, payload, start_date=date.today(), per_contact=per_contact_flag)
            # --- persona_fallback_conf (dal tips) ---
            persona_confs = []
            for m in all_msgs:
                for t in (m.tips or []):
                    if isinstance(t, str) and t.startswith("persona_conf:"):
                        try:
                            persona_confs.append(float(t.split(":",1)[1]))
                        except Exception:
                            pass
            if persona_confs:
                logging["persona_fallback_conf"] = round(sum(persona_confs)/len(persona_confs), 3)
            # bump rules
            plan_inmail_bump_if_no_connection(messages=msgs, actions=acts, calendar=cal, thread=th,
                                              payload=payload, start_date=date.today(), n_days_after_dm=3,
                                              li_status_resolver=None)
            plan_dm_followup_on_connection_accept(messages=msgs, actions=acts, calendar=cal, thread=th,
                                                  payload=payload, start_date=date.today(), n_days_after_dm=3,
                                                  li_status_resolver=None)
            all_msgs += msgs; all_actions += acts; all_calendar += cal
    else:
        default_thread = ThreadSpec(
            name="default",
            persona_id=(payload.buyer_persona_ids or ["generic"])[0],
            channels=["email","linkedin_dm","inmail"],
            cadence=[]
        )
        msgs, acts, cal = compose_thread_sequence(default_thread, payload, start_date=date.today(), per_contact=ALLOW_PER_CONTACT)
        for m in msgs:
            for t in (m.tips or []):
                if t.startswith("persona_conf:"):
                    try: persona_conf_samples.append(float(t.split(":")[1]))
                    except Exception: pass
            m.tips = (m.tips or []) + ["thread:default"]
        plan_inmail_bump_if_no_connection(messages=msgs, actions=acts, calendar=cal, thread=default_thread,
                                          payload=payload, start_date=date.today(), n_days_after_dm=3,
                                          li_status_resolver=None)
        plan_dm_followup_on_connection_accept(messages=msgs, actions=acts, calendar=cal, thread=default_thread,
                                              payload=payload, start_date=date.today(), n_days_after_dm=3,
                                              li_status_resolver=None)
        all_msgs += msgs; all_actions += acts; all_calendar += cal
    # ---- (Opzionale) LLM Polishing su TUTTI i messaggi ----
    if os.getenv("LLM_POLISH_ENABLED","false").lower() in {"1","true","yes"}:
        lang_pol = getattr(payload, "language", "it") or "it"
        for m in all_msgs:
            try:
                if m.text: m.text = _polish_with_llm(m.text, lang=lang_pol)
                if getattr(m, "subject", None): m.subject = _polish_with_llm(m.subject, lang=lang_pol)
            except Exception:
                pass
    # === A/B: filtro opzionale "solo A" o "solo B" ===
    force_variant = (getattr(payload, "force_variant", None) or "").upper()
    if force_variant in {"A", "B"}:
        all_msgs = [m for m in all_msgs if (getattr(m, "variant", None) or "A") == force_variant]
    # === A/B: logging + labels ===
    ab_variants = sorted(set([(getattr(m, "variant", None) or "A") for m in all_msgs]))
    labels_ab = [f"AB:{','.join(ab_variants)}"] if ab_variants else []
    labels = ["Poke-AB"] + labels_ab
    # --- Research block (aggiungi sintesi se c’è) ---
    research_block = None
    if rr_norm:
        if research_summary:
            rr_norm.facts = (rr_norm.facts or [])
            rr_norm.facts.insert(0, EnrichedFact(text=research_summary, fact="Sintesi LLM delle fonti", confidence=0.9))
        research_block = rr_norm
    # --- Rationale safe ---
    rationale_text = "Frameworks compositi; ricerca inline applicata" if (research_hits or research_summary) else "Frameworks compositi; nessuna ricerca"
    # --- Logging (già esistente) ---

    # --- LOGGING/WM -------------------------------------------------
    wm_payload = {
        "contacts": [c.model_dump() for c in (payload.contacts or [])],
        "threads": [t.model_dump() for t in (payload.threads or [])],
        "ts": time.time(),
    }
    wm = _wm_hash(wm_payload)
    logging = {
        "wm_hash": wm,
        "contacts_n": len(payload.contacts or []),
        "threads_n": len(payload.threads or []),
        "ts": time.time(),
        "model": "v3",
        "idempotency_key": (idempotency_key if 'idempotency_key' in locals() else None),
        "ab_variants": (ab_variants if 'ab_variants' in locals() else None),
    }

    if 'persona_synth_id' in locals():
        logging["persona_synth"] = persona_synth_id
    if 'persona_synth_error' in locals():
        logging["persona_synth_error"] = persona_synth_error
    # fallback confidence (dalla tua _closest_persona_id)
    if 'fb_conf' in locals() and fb_conf is not None:
        logging["persona_fallback_confidence"] = float(fb_conf)
    try:
        conf_samples = []
        for m in (all_msgs or []):
            for t in (m.tips or []):
                if isinstance(t, str) and t.startswith("persona_conf:"):
                    conf_samples.append(float(t.split(":",1)[1]))
        if conf_samples:
            logging["persona_fallback_conf"] = round(sum(conf_samples)/len(conf_samples), 3)

    except Exception:
        pass

    what_i_used = WhatIUsed(
        personas=(payload.buyer_persona_ids or None),
        files=["Case studies CPG - Retail.docx"],
        triggers=(payload.triggers.erp if (payload.triggers and payload.triggers.erp) else None)
    )

    coi = COI(status="estimated", note="Ritardi EDI 3–5% → rischio €25–40k/anno",
              assumptions=["Volumi simili a case retail","Ticket medio €20-30"])

    out = StandardOutput(
        messages=all_msgs,
        sequence_next=all_actions,
        calendar=all_calendar,
        labels=labels,
        research=rr_norm,
        logging=log,
        what_i_used=what_i_used,   # <— aggiungi
        coi=coi
    )

    # idempotency store
    if idempotency_key:
        try:
            idem_set(idempotency_key, out)
        except Exception:
            pass

    return out


@app.get("/dm/mapping", response_model=DmMapResponse)
def dm_mapping(industry: Optional[str] = Query(default=None),
               product: Optional[str] = Query(default=None),
               authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    return DmMapResponse(industry=industry, product=product, roles=_dm_mapping(industry, product))

@app.post("/sales/objection_responses", response_model=ObjectionResponses)
def objection_responses(payload: ObjectionInput, authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    out = _price_objection_responses(payload.product, payload.context, payload.currency)
    if os.getenv("LLM_POLISH_ENABLED","false").lower() in {"1","true","yes"}:
        try:
            out = [_polish_with_llm(t, lang=(payload.language or "it")) for t in out]
        except Exception:
            pass
    return ObjectionResponses(responses=out)

@app.post("/threads/simulate", response_model=StandardOutput)
def threads_simulate(payload: GenerateSequenceRequest, authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    start = date.today()
    all_msgs, all_actions, all_cal = [], [], []
    threads = payload.threads or [ThreadSpec(
        name="default",
        persona_id=(payload.buyer_persona_ids or ["generic"])[0],
        channels=["email","linkedin_dm","inmail"], cadence=[]
    )]
    for th in threads:
        # rispetta toggles per_contact, bump/DM ecc.
        msgs, acts, cal = compose_thread_sequence(th, payload, start_date=start,
                                                  per_contact=bool(getattr(th, "meta",{}).get("per_contact", False)))
        # solo simulazione: nessun salvataggio
        all_msgs += msgs; all_actions += acts; all_cal += cal
        # --- OPTIONAL: LLM polishing (sotto feature-flag) ---
        if os.getenv("LLM_POLISH_ENABLED","false").lower() in {"1","true","yes"}:
            lang_pol = getattr(payload, "language", "it") or "it"
            for m in all_msgs:
                try:
                    m.text = _polish_with_llm(m.text or "", lang=lang_pol)
                    if getattr(m, "subject", None):
                        # raffinatura leggera dell'oggetto (riusa la stessa funzione)
                        m.subject = _polish_with_llm(m.subject or "", lang=lang_pol)
                except Exception:
                    # non bloccare il flusso
                    pass

    # facoltativo: genera ICS inline
    ics = _ics_from_calendar(all_cal)
    return StandardOutput(messages=all_msgs, sequence_next=all_actions, calendar=all_cal,
                          labels=["simulation"], what_i_used=None, rationale="simulate", logging={"ics": ics})

@app.post("/calendar/export_ics")
def calendar_export_ics(payload: StandardOutput, authorization: Optional[str] = Header(None)):
    """
    Accetta direttamente lo StandardOutput (quello restituito da /run/generate_sequence)
    e produce un .ics con gli eventi in 'calendar'.
    """
    ensure_auth(authorization)
    ics_text = _ics_from_calendar(payload.calendar or [], cal_name="Flowagent v3 Sequence")
    return Response(
        content=ics_text,
        media_type="text/calendar",
        headers={"Content-Disposition": 'attachment; filename="flowagent_sequence.ics"'}
    )

@app.post("/run/rank", response_model=RankResponse)
def rank(payload: RankRequest, authorization: Optional[str] = Header(None)):
    _guard_ip(json.dumps(payload.dict() if hasattr(payload, 'dict') else payload))
    ensure_auth(authorization)
    ranked = [{"id": it.id, "score": round(0.9 - i*0.07, 3), "reason": f"Coerenza {payload.objective}"} for i,it in enumerate(payload.items)]
    return RankResponse(ranked=ranked)

TRIGGER_RULES = {
    "industry": {"retail": ["omnichannel","shelf-availability","e-invoicing"]},
    "role": {"cio": ["integration-lag","cost-to-serve"], "operations": ["onboarding-lead-time"]},
    "erp": {"sap": ["s4hana-migration","idoc-mapping"], "oracle": ["fusion-integration"]},
}

@app.post("/sequence/triggers", response_model=TriggerExtractResp)
def extract_triggers(payload: GenerateSequenceRequest, authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    # a) KB + web (se attivo)
    ropt = _effective_research_mode(payload.research or ResearchOptions())
    kb = _do_kb_search(_company_query(payload), top_k=5)
    web = _perform_research(payload) if ropt.mode == "web" else []

    corpus = " ".join([(x.get("snippet") or x.get("title") or "") for x in (kb + web)])
    trg = {
        "industry": _top_matches(corpus, ["retail", "CPG", "automotive", "utilities"]),
        "company": _top_matches(corpus, ["acquisizione", "partnership", "apertura hub", "nuovo ERP"]),
        "role": _top_matches(corpus, ["CIO", "Head of Integration", "Procurement", "Logistics"]),
        "personal": [],  # puoi popolare da payload.contacts[*].linkedin_profile se disponibile
        "erp": _top_matches(corpus, ["SAP", "Oracle", "Microsoft Dynamics", "JD Edwards"])
    }
    # b) ranking combinazioni semplici (punteggi pesati)
    combos = []
    for ind in trg["industry"][:2]:
        for erp in trg["erp"][:2]:
            score = 1.0 + 0.7*corpus.count(ind.lower()) + 0.6*corpus.count(erp.lower())
            combos.append({"combo": {"industry": ind, "erp": erp}, "score": score})
    combos.sort(key=lambda x: x["score"], reverse=True)
    return TriggerExtractResp(triggers=trg, ranked=combos)

@app.post("/validate/compliance", response_model=ComplianceResponse)
def compliance_validate(
    payload: ComplianceRequest,
    authorization: Optional[str] = Header(None),
):
    ensure_auth(authorization)

    # Normalizzazioni safe
    text = payload.text or ""
    text_low = text.lower()
    rules = payload.rules or {}
    if not isinstance(rules, dict):
        rules = {}

    # Estrai e pulisci banned_terms
    raw_banned = rules.get("banned_terms", [])
    if not isinstance(raw_banned, list):
        raw_banned = [raw_banned]
    banned_terms = []
    for x in raw_banned:
        try:
            banned_terms.append(str(x))
        except Exception:
            continue

    require_cta = bool(rules.get("require_cta", True))
    anti_jargon = bool(rules.get("anti_jargon", True))
    max_words   = rules.get("max_words", None)

    violations = []

    # 1) CTA presente?
    if require_cta:
        tokens = ("?", "prenot", "disponibile", "chi")
        if not any(tok in text_low for tok in tokens):
            violations.append({"code": "CTA_MISSING", "detail": "Manca una call-to-action chiara"})

    # 2) Anti-jargon
    if anti_jargon and banned_terms:
        for bad in banned_terms:
            if bad.lower() in text_low:
                violations.append({"code": "JARGON", "detail": f"Termine vietato: {bad}"})
                break

    # 3) Lunghezza massima parole
    if max_words is not None:
        try:
            mw = int(max_words)
            words = len(text.split())
            if words > mw:
                violations.append({"code": "TOO_LONG", "detail": f"Testo {words} parole > max {mw}"})
        except Exception:
            # se max_words non è castabile, ignora senza crash
            pass

    return ComplianceResponse(**{"pass": len(violations)==0, "violations": violations})

@app.post("/calendar/build", response_model=CalendarBuildResponse)
def calendar_build(
    payload: CalendarBuildRequest,
    authorization: Optional[str] = Header(None),
):
    ensure_auth(authorization)

    events = []
    start = payload.start_date if payload.start_date else date.today()
    rules = payload.rules or CalendarRules()

    base = payload.base_sequence or [SequenceAction(day=0, action="Step 1")]
    for sa in base:
        d = start + timedelta(days=sa.day)
        if rules.no_weekend:
            d = next_workday(d)
        events.append(
            CalendarEvent(date=d, action=sa.action, no_weekend_respected=True)

        )

    # (eventuale logica su payload.signals...)

    ics = "BEGIN:VCALENDAR\nVERSION:2.0\nPRODID:-//FlowagentV3//EN\nEND:VCALENDAR"
    return CalendarBuildResponse(calendar=events, ics=ics)

@app.post("/kb/ingest_advanced", response_model=KBIngestResponse)
async def kb_ingest_advanced(req: KBIngestRequest):
    """
    Ingest “ricco”: PDF/images -> testo, estrae URL e salva .txt di fianco al file originale.
    Se `content_base64` è presente salva su kb/raw; se `url` punta a S3/HTTP, scarica best-effort.
    """
    # 1) salva file
    dest_dir = UPLOAD_DIR
    dest_dir.mkdir(parents=True, exist_ok=True)

    if req.content_base64:
        raw = base64.b64decode(req.content_base64)
        fname = req.metadata.get("filename","upload.bin") if req.metadata else "upload.bin"
        path = dest_dir / fname
        path.write_bytes(raw)
    elif req.url:
        # se preferisci non scaricare lato server, rimuovi questo ramo:
        import requests as _rq
        r = _rq.get(req.url, timeout=DEFAULT_TIMEOUT)
        r.raise_for_status()
        fname = Path(req.url.split("?")[0]).name or "remote.bin"
        path = dest_dir / fname
        path.write_bytes(r.content)
    else:
        raise HTTPException(400, "Missing content_base64 or url")

    # 2) estrai testo
    text = ""
    suf = path.suffix.lower()
    if suf == ".pdf":
        text = _extract_text_from_pdf(str(path))
    elif suf in (".png",".jpg",".jpeg",".tif",".tiff"):
        text = _extract_text_from_image(str(path))
    elif suf in (".txt",".md"):
        text = path.read_text(encoding="utf-8", errors="ignore")
    elif suf == ".docx":
        try:
            import docx2txt
            text = docx2txt.process(str(path)) or ""
        except Exception:
            text = ""
    # 3) salva sidecar .txt + raccogli URL
    url_list = _extract_urls(text)
    (path.with_suffix(".txt")).write_text(text, encoding="utf-8", errors="ignore")
    from research_agent import research_extract
    if url_list:
        try:
            research_extract("\n".join(url_list))
        except Exception:
            pass
    return KBIngestResponse(doc_id=path.name, chunks=(1 if text else 0))

@app.get("/kb/personas/index", response_model=PersonaIndexResponse)
def kb_personas_index(
    authorization: Optional[str] = Header(None),
    raw: Optional[bool] = Query(default=False, description="Se true e PERSONA_INDEX_JSON esiste, ritorna il JSON grezzo")
):
    ensure_auth(authorization)

    # se vuoi esporre direttamente un file indice già pronto
    if raw and os.path.exists(PERSONA_INDEX_JSON):
        try:
            blob = Path(PERSONA_INDEX_JSON).read_text(encoding="utf-8")
            return Response(content=blob, media_type="application/json")
        except Exception:
            pass  # se fallisce, ricadi sull’indice ricostruito

    # ricarica indice runtime dai .docx e dalle synth-* se le includi nel loader
    docs = _load_persona_index(force=True)  # force=True per avere stato fresco
    items: List[PersonaIndexItem] = []
    for d in docs:
        items.append(PersonaIndexItem(
            persona_id=d.persona_id,
            industry=d.industry,
            role=d.role,
            pains=len(d.pains or []),
            kpis=len(d.kpis or []),
            objections=len(d.objections or []),
        ))
    return PersonaIndexResponse(items=sorted(items, key=lambda x: x.persona_id.lower()))

@app.get("/kb/personas/show")
def kb_persona_show(
    id: str = Query(..., description="persona_id o nome file (senza estensione)"),
    authorization: Optional[str] = Header(None)
):
    ensure_auth(authorization)
    base_dir = os.getenv("KB_DIR", "./kb")
    personas_dir = os.path.join(base_dir, "personas")
    if not os.path.isdir(personas_dir):
        raise HTTPException(status_code=404, detail="Personas library non trovata")

    target = (id or "").strip()
    if not target:
        raise HTTPException(status_code=400, detail="id mancante")
    target_cf = target.casefold()

    candidates: List[str] = []

    # Prova 1: match diretto su file .json (case-insensitive, match parziale in coda)
    for fn in os.listdir(personas_dir):
        if not fn.lower().endswith(".json"):
            continue
        name = os.path.splitext(fn)[0]
        if name.casefold() == target_cf or name.casefold().endswith(target_cf):
            candidates.append(os.path.join(personas_dir, fn))

    # Prova 2: indice PERSONA_INDEX_JSON → mappa persona_id → file_name
    if not candidates:
        index_path = os.path.join(personas_dir, "PERSONA_INDEX_JSON")
        if os.path.isfile(index_path):
            try:
                idx = json.loads(Path(index_path).read_text(encoding="utf-8"))
                for p in (idx.get("personas") or []):
                    pid = (p.get("persona_id") or "").casefold()
                    if pid == target_cf:
                        fn = p.get("file_name") or f"{id}.json"
                        path = os.path.join(personas_dir, fn)
                        if os.path.isfile(path):
                            candidates.append(path)
                            break
            except Exception:
                pass

    if not candidates:
        raise HTTPException(status_code=404, detail=f"Persona '{id}' non trovata")

    path = candidates[0]
    try:
        obj = json.loads(Path(path).read_text(encoding="utf-8"))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Errore lettura persona: {e}")

    return JSONResponse(content=obj)
# === END PATCH ===


@app.post("/kb/ingest", response_model=KBIngestResponse)
def kb_ingest(
    payload: KBIngestRequest,
    authorization: Optional[str] = Header(None),
    approved_hdr: Optional[str] = Header(default=None, alias=APPROVAL_HEADER),
    approved_q: Optional[str] = Query(default=None, alias="approved"),   # <-- NEW
):
    _guard_ip_from_request(payload)
    ensure_auth(authorization)
    approval_gate(True, approved_hdr or approved_q)

    # 1) salva il file o la sorgente (dipende dal tuo KBIngestRequest)
    #    Qui assumo KBIngestRequest abbia: content (base64 o testo) e filename
    kb_root = Path("kb/raw"); kb_root.mkdir(parents=True, exist_ok=True)
    fname = payload.filename or f"doc_{int(time.time())}.txt"
    fp = kb_root / fname
    try:
        if payload.content_base64:
            import base64
            fp.write_bytes(base64.b64decode(payload.content_base64))
        elif payload.text:
            fp.write_text(payload.text, encoding="utf-8")
        else:
            raise ValueError("payload vuoto")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Ingest fallita: {e}")

    # 2) genera sidecar .txt se era docx/pdf/immagine (placeholder semplice)
    side = fp.with_suffix(".txt")
    try:
        if fp.suffix.lower() == ".docx":
            import docx
            doc = docx.Document(str(fp))
            txt = "\n".join([p.text for p in doc.paragraphs])
            side.write_text(txt, encoding="utf-8")
        elif fp.suffix.lower() in {".txt",".md"}:
            side.write_text(fp.read_text(encoding="utf-8", errors="ignore"), encoding="utf-8")
        else:
            # per pdf/imaging servirebbe OCR: placeholder minimo
            side.write_text(fp.name, encoding="utf-8")
    except Exception:
        pass

    # 3) ritorna conteggio “chunk” semplificato e doc_id
    txt = side.read_text(encoding="utf-8", errors="ignore") if side.exists() else ""
    chunks = max(1, len(txt)//800)
    return KBIngestResponse(doc_id=fp.name, chunks=chunks)


@app.get("/kb/search", response_model=KBSearchResponse)
def kb_search(
    q: str,
    industry: Optional[str] = None,
    role: Optional[str] = None,
    lang: Optional[str] = None,
    top_k: int = 5,
    authorization: Optional[str] = Header(None),
):
    ensure_auth(authorization)
    # indicizza KB "flat"
    items = get_kb_items_flat()

    # filtri opzionali su title/snippet
    if any([industry, role, lang]):
        fitems = []
        for it in items:
            blob = " ".join([it.get("title",""), it.get("snippet","")]).lower()
            if industry and industry.lower() not in blob: continue
            if role and role.lower() not in blob: continue
            if lang and lang.lower() not in blob: continue
            fitems.append(it)
        items = fitems

    # preferisci TF-IDF; se non disponibile, fallback al count-based
    try:
        hits = _do_kb_search_tfidf(q, items, top_k=top_k) or []
        if not hits:
            hits = _do_kb_search_countbased(q, items, top_k=top_k) or []
    except Exception:
        hits = _do_kb_search_countbased(q, items, top_k=top_k) or []

    return KBSearchResponse(
        query=q,
        matches=[
            KBSearchMatch(file=h.get("url") or h.get("title",""),
                          score=float(h.get("score_sim") or h.get("score_count") or 0.0),
                          snippet=h.get("snippet"))
            for h in hits
        ],
    )

@app.get("/kb/list", response_model=KBListResponse)
def kb_list(industry: Optional[str]=None, role: Optional[str]=None, lang: Optional[str]=None,
            q: Optional[str]=None, authorization: Optional[str] = Header(None), as_zip: Optional[bool] = Query(default=False)):
    ensure_auth(authorization)
    items=[]
    for p in Path("kb/raw").glob("**/*"):
        if p.is_file() and p.suffix.lower() in (".pdf",".png",".jpg",".jpeg",".tif",".tiff",".docx",".txt",".md"):
            meta={"size": p.stat().st_size, "ext": p.suffix.lower()}
            if q:
                # cerca nel sidecar .txt se esiste
                try:
                    txt = (p.with_suffix(".txt")).read_text(encoding="utf-8", errors="ignore")
                except Exception:
                    txt = ""
                if q.lower() not in (p.name.lower()+" "+txt.lower()):
                    continue
            items.append(KBListItem(doc_id=p.name, title=p.stem, lang=lang, metadata=meta))
    docs=[]
    for p in Path("kb/raw").glob("**/*"):
        if p.is_file():
            docs.append(KBListItem(doc_id=p.name, title=p.stem, lang=None, metadata={"size": p.stat().st_size}))

    kb_root = Path("kb/raw")
    docs = []
    if kb_root.exists():
        for f in kb_root.glob("*.docx"):
            docs.append(KBListItem(doc_id=f.name, title=f.stem, lang="it", metadata={"path": str(f)}))

    if not as_zip:
        return KBListResponse(docs=docs)

    # prepara ZIP in memoria
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for item in docs:
            p = Path(item.metadata["path"])
            try:
                z.write(p, arcname=p.name)
            except Exception:
                continue
    buf.seek(0)
    headers = {
        "Content-Disposition": 'attachment; filename="kb_raw.zip"',
        "X-Total-Files": str(len(docs)),
    }
    return StreamingResponse(buf, media_type="application/zip", headers=headers)

@app.delete("/kb/delete")
def kb_delete(doc_id: str, soft: bool = True, authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    p = Path("kb/raw")/doc_id
    if not p.exists():
        raise HTTPException(404, "not found")
    if soft:
        # soft delete = rinomina
        p.rename(p.with_suffix(p.suffix+".deleted"))
        return {"ok": True, "soft": True}
    else:
        p.unlink(missing_ok=True)
        (p.with_suffix(".txt")).unlink(missing_ok=True)
        return {"ok": True, "soft": False}
@app.get("/kb/doc/{doc_id}")
def kb_doc_get(doc_id: str):
    doc = kb_get(doc_id)  # implementa secondo il tuo storage
    if not doc or doc.get("private"):
        raise HTTPException(404, "Not found")
    snippet = doc["content"][:preview_chars]
    return {
        "doc_id": doc_id,
        "title": doc.get("title"),
        "lang": doc.get("lang",""),
        "snippet": snippet,
        "metadata": doc.get("metadata",{}),
    }

@app.post("/company/evidence/upsert", response_model=UpsertResponse)
def company_evidence_upsert(
    payload: CompanyEvidence,
    authorization: Optional[str] = Header(None),
    approved_hdr: Optional[str] = Header(default=None, alias=APPROVAL_HEADER),
    approved_q: Optional[str] = Query(default=None, alias="approved"),
):
    ensure_auth(authorization)
    approval_gate(True, approved_hdr or approved_q)

    labels = []
    for e in (payload.erp or []):
        labels.append(f"CompanyEvidence: {e}")

    stored = payload.model_dump() if hasattr(payload, "model_dump") else payload.dict()
    return UpsertResponse(ok=True, stored=stored)

# ===================== BEGIN PATCH: /competitor/analysis fix =====================
@app.post("/competitor/analysis", response_model=CompetitorMatrix)
def competitor_analysis(payload: CompetitorRequest, authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    ropt = _effective_research_mode(payload.research or ResearchOptions())
    sources: list[dict] = _perform_research(payload) if ropt.mode != "off" else []

    # 1) split corpus per competitor (dai risultati web)
    corpus_map = _build_corpus_per_competitor(sources, payload.competitors or [])

    # 2) normalizza features dalla richiesta (NON _normalized_text)
    normalized_features = _norm_features(payload)

    # 3) coverage (0..1) per (feature × competitor)
    feature_coverage = {f: {} for f in normalized_features}
    for comp in (payload.competitors or []):
        text_corpus = (corpus_map.get(comp.lower(), "") or "")
        for feat in normalized_features:
            base  = _score_feature_coverage(feat, corpus=text_corpus)
            bonus = 0.25 * _feature_hit_score(feat, text_corpus)   # <— QUI IL BONUS
            feature_coverage[feat][comp] = min(1.0, base + bonus)

    # 4) ranking totale (pesi industry/icp + coverage)
    ranking = _rank_competitors(
        feature_coverage=feature_coverage,
        industry=payload.industry,
        icp=payload.icp,
        corpus_map=corpus_map
    )

    # 5) items con score_* (feature/industry_icp/total)
    items: list[CompetitorItem] = []
    for comp in (payload.competitors or []):
        cov = {feat: feature_coverage[feat].get(comp, 0.0) for feat in normalized_features}
        item = CompetitorItem(name=comp, feature_coverage=cov)
        item = _score_competitor_item(
            item, normalized_features, corpus_map.get(comp.lower(), ""), payload.industry, payload.icp
        )
        items.append(item)

    # 6) opzionale: summary LLM
    research_summary = None
    if os.getenv("LLM_SUMMARY_ENABLED","false").lower() in {"1","true","yes"} and sources:
        try:
            research_summary = _summarize_corpus_items(sources, lang=(payload.language or "it"))
        except Exception:
            research_summary = None

    return CompetitorMatrix(
        company=payload.company,
        target=payload.company,
        industry=payload.industry,
        icp=payload.icp,
        normalized_features=normalized_features,
        feature_coverage=feature_coverage,
        items=items,
        research_summary=research_summary,
        scoring={"ranking": ranking}
    )
# ===================== END PATCH: /competitor/analysis fix =====================



@app.get("/analysis/competitors/table")
async def competitors_table_csv(
    company: str = Query(description="Azienda target"),
    competitors: str = Query(description="Lista competitor, separati da virgola"),
    industry: Optional[str] = Query(None, description="Industry/settore target"),
    icp: Optional[str] = Query(None, description="ICP opzionale"),
    language: Optional[str] = Query("it", description="Lingua"),
    research_mode: Optional[str] = Query("kb_only", description="off|kb_only|web"),
    max_sources: Optional[int] = Query(8, ge=1, le=50, description="limite fonti ricerca"),
    authorization: Optional[str] = Header(None),
):
    """
    Restituisce direttamente un CSV (feature × competitor) pronto per Excel.
    Usa build_competitor_matrix(...) e normalizza la tabella.
    """
    ensure_auth(authorization)

    comp_list = [c.strip() for c in (competitors or "").split(",") if c.strip()]

    payload = CompetitorRequest(
        company=company, competitors=comp_list, industry=industry, icp=icp, language=language,
        research=ResearchOptions(mode=research_mode, max_sources=max_sources)
    )
    # build (compatibile con async/sync)
    matrix_resp = competitor_analysis(payload, authorization)  # riusa la funzione sopra
    rows = _matrix_to_feature_table_rows(matrix_resp.matrix)
    csv_text = _rows_to_csv(rows)
    safe_company = re.sub(r"[^A-Za-z0-9_-]+", "_", company).strip("_") or "company"
    filename = f"competitor_matrix_{safe_company}.csv"
    return StreamingResponse(io.BytesIO(csv_text.encode("utf-8")),
                             media_type="text/csv",
                             headers={"Content-Disposition": f'attachment; filename="{filename}"'})

@app.get("/company/evidence/{company_id}", response_model=CompanyEvidence)
def company_evidence_get(
    company_id: str,
    authorization: Optional[str] = Header(None),
):
    ensure_auth(authorization)
    return CompanyEvidence(company_id=company_id, erp=["SAP"], competitor=["X"], tools=["Y"])

@app.post("/signals/record", response_model=UpsertResponse)
def record_signal(
    payload: ManualSignal,
    authorization: Optional[str] = Header(None),
    approved_hdr: Optional[str] = Header(default=None, alias=APPROVAL_HEADER),
    approved_q: Optional[str] = Query(default=None, alias="approved"),
):
    ensure_auth(authorization)
    approval_gate(True, approved_hdr or approved_q)      # <- azione mutante => richiede conferma in prod

    stored = payload.model_dump() if hasattr(payload, "model_dump") else payload.dict()
    return UpsertResponse(ok=True, stored=stored)

@app.post("/coi/estimate", response_model=COIEstimateResponse)
def coi_estimate(payload: COIEstimateRequest, authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    return COIEstimateResponse(
        status="estimated",
        note="COI ~ €25–40k/anno",
        assumptions=["5% errori","AOV €120","Process cost €2/ordine"]
    )

@app.post("/ab/promote", response_model=ABPromoteResponse)
def ab_promote(
    payload: ABPromoteRequest,
    authorization: Optional[str] = Header(None),
    approved_hdr: Optional[str] = Header(default=None, alias=APPROVAL_HEADER),
    approved_q: Optional[str] = Query(default=None, alias="approved"),
):
    _guard_ip(json.dumps(payload.dict() if hasattr(payload, 'dict') else payload))
    ensure_auth(authorization)
    approval_gate(True, approved_hdr or approved_q)
    data = payload.model_dump() if hasattr(payload, "model_dump") else payload.dict()
    return ABPromoteResponse(ok=True, sequence_id=data["sequence_id"], variant_id=data["variant_id"])

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None  # se manca la libreria, gestiamo sotto

@app.post("/export/docx")
def export_docx(payload: StandardOutput, authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    blob = _docx_from_messages(payload.messages or [], title="Flowagent Sequence")
    if not blob:
        raise HTTPException(status_code=501, detail="DOCX export non disponibile: installa 'python-docx'.")
    return Response(
        content=blob,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="flowagent_sequence.docx"'}
    )

@app.get("/export/sequence/{sequence_id}", response_model=ExportSequenceResponse)
def export_sequence(sequence_id: str, format: str = "json", authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    if format == "csv":
        return ExportSequenceResponse(ok=True, format="csv", sequence_id=sequence_id)
    return ExportSequenceResponse(ok=True, format="json", sequence_id=sequence_id)

@app.post("/export/sequence/csv")
def export_sequence_csv(out: StandardOutput):
    def _iter():
        yield "channel,step,subject,text\n"
        for m in (out.messages or []):
            row = [
                (m.channel or ""),
                (m.step or ""),
                (m.subject or "").replace("\n"," ").replace(",",";"),
                (m.text or "").replace("\n"," ").replace(",",";"),
            ]
            yield ",".join(row) + "\n"
    return StreamingResponse(_iter(), media_type="text/csv", headers={"Content-Disposition":"attachment; filename=sequence.csv"})

@app.post("/export/sequence/docx")
def export_sequence_docx(payload: GenerateSequenceRequest, authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    thread = (payload.threads or [ThreadSpec()])[0]
    msgs, acts, cal, log = compose_thread_sequence(thread, payload, start_date=date.today(), per_contact=False)

    _ = (acts, cal, log)

    doc = Document()
    doc.add_heading("FlowAgent – Sequenza", 0)
    for m in msgs:
        h = doc.add_heading(f"{m.channel.upper()} | {m.step} | {m.variant}", level=1)
        if m.subject: doc.add_paragraph(f"Subject: {m.subject}")
        doc.add_paragraph(m.text or "")
        if m.tips:
            doc.add_paragraph("TIPPS:")
            for t in m.tips: doc.add_paragraph(f"• {t}", style="List Bullet")
        doc.add_paragraph("")

    bio = BytesIO()
    doc.save(bio); bio.seek(0)
    return StreamingResponse(bio, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": 'attachment; filename="sequence.docx"'})

@app.post("/calendar/ics", response_model=CalendarBuildResponse)
def calendar_ics(req: CalendarBuildRequest, authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    start = req.start_date or date.today()
    rules = req.rules or CalendarRules()
    cal_events = []
    for a in (req.base_sequence or []):
        d = _business_day_add(start, a.day, rules)
        cal_events.append(CalendarEvent(date=d, action=a.action, no_weekend_respected=rules.no_weekend))
    ics_text = _ics_from_calendar(cal_events, cal_name="FlowAgent v3")
    return CalendarBuildResponse(calendar=cal_events, ics=ics_text)

@app.post("/export/ics")
def export_ics(payload: StandardOutput, authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    ics = _ics_from_calendar(payload.calendar or [])
    return HTMLResponse(content=ics, media_type="text/calendar")

@app.get("/reminders/today", response_model=DailyReminders)
def reminders_today(tz: Optional[str] = "Europe/Rome", authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    # in reale: leggere da storage le calendar actions generate ieri/oggi
    today = date.today().isoformat()
    # mock live: filtra da un archivio locale se presente
    return DailyReminders(date=today, items=[
        {"time": "09:00", "action": "Email Step 1 → Mario Rossi"},
        {"time": "11:00", "action": "LinkedIn DM → Sara Bianchi"},
    ])

@app.get("/export/campaigns.csv")
def export_campaigns_csv(authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)
    rows = _load_campaigns()
    buf = io.StringIO()
    fieldnames = sorted(set(k for r in rows for k in r.keys())) or ["date","contact","action"]
    w = csv.DictWriter(buf, fieldnames=fieldnames)
    w.writeheader()
    for r in rows:
        w.writerow(r)
    return Response(content=buf.getvalue(), media_type="text/csv",
                    headers={"Content-Disposition": 'attachment; filename="campaigns.csv"'})

@app.get("/export/library.zip")
def export_library_zip():
    import zipfile, io
    mem = io.BytesIO()
    with zipfile.ZipFile(mem, mode="w") as z:
        base = Path("kb/raw")
        for p in base.glob("*.docx"):
            z.write(p, arcname=p.name)
    mem.seek(0)
    return Response(content=mem.read(), media_type="application/zip", headers={"Content-Disposition":"attachment; filename=library.zip"})

@app.get("/kb/libraries/zip")
def kb_libraries_zip():
    """Zippa tutte le librerie in /kb (personas, frameworks, cases)."""
    base = Path("kb")
    zip_path = base / "libraries.zip"
    with zipfile.ZipFile(zip_path, "w") as zf:
        for f in base.rglob("*.json"):
            zf.write(f, f.relative_to(base))
    return FileResponse(str(zip_path), filename="libraries.zip")

# --- JSON ---
@app.post("/export/messages.json")
def export_json(payload: StandardOutput):
    return payload.model_dump()

# --- CSV ---
import csv, io
@app.post("/export/messages.csv")
def export_csv(payload: StandardOutput):
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["thread","channel","step","variant","subject","text"])
    for m in (payload.messages or []):
        w.writerow([getattr(m,"thread", ""), m.channel, m.step, m.variant, getattr(m,"subject",""), m.text])
    return HTMLResponse(content=buf.getvalue(), media_type="text/csv")

# --- DOCX (Word) ---
@app.post("/export/messages.docx")
def export_docx(payload: StandardOutput):
    doc = Document()
    for m in (payload.messages or []):
        doc.add_heading(f"{m.channel} · step {m.step}", level=3)
        if getattr(m,"subject",None):
            doc.add_paragraph(f"Subject: {m.subject}")
        doc.add_paragraph(m.text)
        doc.add_paragraph("---")
    tmp = "/tmp/messages.docx"
    doc.save(tmp)
    with open(tmp, "rb") as f:
        data = f.read()
    return HTMLResponse(content=data, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.post("/send/email", response_model=SendEmailResponse)
def send_email(
    payload: SendEmailRequest,
    authorization: Optional[str] = Header(None),
    approved_hdr: Optional[str] = Header(default=None, alias=APPROVAL_HEADER),
    # Query: ?approved=true|yes|1|on
    approved_q: Optional[str] = Query(default=None, alias="approved"),
):
    _guard_ip(json.dumps(payload.dict() if hasattr(payload, 'dict') else payload))
    ensure_auth(authorization)
    approval_gate(True, approved_hdr or approved_q)

    msg_id = f"msg_{int(datetime.now(timezone.utc).timestamp())}"
    queued = datetime.now(timezone.utc)
    return SendEmailResponse(message_id=msg_id, provider=payload.provider, queued_at=queued)

WEBHOOK_SECRET = os.getenv("WEBHOOK_SECRET", "")

@app.post("/webhooks/events", include_in_schema=False)
def webhooks (

    x_webhook_secret: Optional[str] = Header(None, alias="X-Webhook-Secret"),
):
    if not WEBHOOK_SECRET or x_webhook_secret != WEBHOOK_SECRET:
        raise HTTPException(status_code=403, detail="Invalid webhook secret")

    # TODO: gestisci evento
    return {"ok": True}

# --- PATCH: web_search live ---
@app.post("/research/web_search")
def research_web_search(
    payload: dict = Body(...),
    authorization: Optional[str] = Header(None),
):
    ensure_auth(authorization)
    # delega vera al research agent
    items = []
    try:
        from research_agent import _perform_research
        raw = _perform_research(payload) or []
        # app.py ha già _normalize(items) che produce ResearchResult
        norm = _normalize(raw)
        for c in norm.citations:
            items.append({
                "url": c.url,
                "title": c.title or "",
                "snippet": next((f.text for f in norm.facts if f.source == c.url), ""),
                "published_at": datetime.now(timezone.utc).isoformat()
            })
    except Exception:
        # fallback: niente crash in prod
        items = []
    max_results = min(int(payload.get("max_results", 5)), 10)
    return {"results": items[:max_results]}

# --- PATCH: extract live (bozza robusta) ---
@app.post("/research/extract")
def research_extract(
    payload: dict = Body(...),
    authorization: Optional[str] = Header(None),
):
    ensure_auth(authorization)
    url = (payload or {}).get("url") or ""
    if not url:
        raise HTTPException(status_code=400, detail="url required")
    try:
        from research_agent import _extract_url  # se l'hai; altrimenti usa il tuo fetcher
        text = (payload or {}).get("text") or ""
        data = _extract_url(url)  # {title, summary, key_points}
        res = research_extract(text)
        return {"urls": res["urls"], "fetched": res["fetched"]}
    except Exception:
        return {"url": url, "title": "", "summary": "", "key_points": []}

# --- PATCH: enrich_contacts live (bozza robusta) ---
@app.post("/research/enrich_contacts")
def research_enrich_contacts(
    payload: dict = Body(...),
    authorization: Optional[str] = Header(None),
):
    ensure_auth(authorization)
    contacts = (payload or {}).get("contacts") or []
    out = []
    try:
        from research_agent import research_enrich_contacts  # se disponibile
        out = research_enrich_contacts(contacts) or []
    except Exception:
        out = []
    return {"enriched": out}

@app.post("/content/linkedin_creatives", response_model=LinkedInCreativesRes)
def content_linkedin_creatives(payload: LinkedInCreativesReq,
                               authorization: Optional[str] = Header(None)):
    ensure_auth(authorization)

    # 1) recupera persona base
    persona = payload.base_persona or {}
    if not persona and payload.persona_id and callable(globals().get("get_persona_by_id", None)):
        persona = get_persona_by_id(payload.persona_id) or {}

    # 2) synth creative (se richiesto)
    if payload.creative_enabled and callable(globals().get("synthesize_creative_persona", None)):
        persona_creative = synthesize_creative_persona(persona, sector=(payload.sector or "generic"),
                                                       language=(payload.language or "it"))
    else:
        persona_creative = {"enabled": False}

    # 3) genera A/B
    persona_id = (payload.persona_id or "").strip()
    if not persona_id:
        tmp = _get_or_synth_persona(payload.industry or "generic", payload.role or "generic", lang=payload.lang or "it")
        persona_id = tmp["id"]
    ab = generate_linkedin_creatives(theme=payload.theme,
                                     persona_creative=persona_creative,
                                     evidence=payload.evidence,
                                     language=(payload.language or "it"),
                                     coi_enabled=bool(payload.coi_enabled))
    return LinkedInCreativesRes(**ab)

if os.getenv("SELFTEST_MODELS","false").lower() in {"1","true","yes"}:
    # istanziazione minima per verificare serializzazione
    _ = Message(channel="email", step="step-00", variant="A", subject=None, text="ciao", tips=[])
    _ = SequenceAction(channel="email", step="step-00", action="send", date=date.today())
    _ = CalendarEvent(date=str(date.today()), action="send", channel="email", step="step-00")
    _ = ResearchOptions()
    _ = GenerateSequenceRequest(threads=[], contacts=[], buyer_persona_ids=[])
    _ = RankRequest(candidates=[{"id":"x"}]); _ = RankResponse(items=[RankItem(id="x", score=0.9, payload={})])
    _ = CalendarBuildResponse(calendar=[], ics=None)
    _ = COIEstimateRequest(signals=[ManualSignal(name="x", value=1)], horizon_days=30)
    _ = COIEstimateResponse(status="ok", estimate=123.4, notes=["ok"])
    _ = KBSearchMatch(file="a.txt", score=0.7); _ = KBListItem(file="a.txt", kind="txt", size=10)
    _ = UpsertResponse(ok=True, id="1")
    _ = CompetitorRequest(company="ACME", competitors=["Foo","Bar"])

#todo
#todo---------@app.get("/_health/models")
#---------def _health_models():
    # istanze minime
#_ = ComplianceResponse(True, violations=[])
#_ = CalendarBuildRequest(start_date=date.today(), base_sequence=[])
# _ = CalendarBuildResponse(calendar=[], ics="BEGIN:VCALENDAR\nEND:VCALENDAR")
# _ = KBSearchResponse(query="test", matches=[KBSearchMatch(file="a.docx", score=0.8, snippet="...")])
#   _ = UpsertResponse(ok=True, id="123", stored={"x":1})
#    _ = ExportSequenceResponse(ok=True, format="json", sequence_id="abc")
#   _ = CompetitorRequest(company="ACME", competitors=["X","Y"], research=ResearchOptions())
#    _ = ResearchResult(facts=[EnrichedFact(text="ok")], citations=[])
#    _ = StandardOutput(messages=[], coi=COI(), sequence_next=[], calendar=[])
#    return {"ok": True}

